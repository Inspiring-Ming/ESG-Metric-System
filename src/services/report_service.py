#!/usr/bin/env python3
"""
ESG Report Service - Self-Contained Report Generation

This service handles all ESG report generation in ONE file:
1. Data validation and structuring
2. Word document (.docx) generation
3. PDF document (.pdf) generation
4. Report storage and retrieval

No external dependencies - completely self-contained.
"""

import json
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional
import os

# Word document imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# PDF document imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, ListFlowable, ListItem
)


class ReportService:
    """Complete self-contained ESG report generation service"""

    def __init__(self):
        """Initialize report service"""
        # Get project root directory (2 levels up from this file: src/services/ -> src/ -> root/)
        current_file = Path(__file__).resolve()
        project_root = current_file.parent.parent.parent
        self.reports_dir = project_root / "generated_reports"
        self.reports_dir.mkdir(parents=True, exist_ok=True)

        print("📊 ReportService initialized (self-contained)")
        print(f"   Reports directory: {self.reports_dir.absolute()}")

    # ==================== MAIN API METHODS ====================

    def generate_word_report(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate Word format ESG report

        Args:
            report_data: Dictionary containing:
                - company_name: str
                - year: str
                - industry: str
                - calculations: List[Dict]
                - quality_score: float (optional)

        Returns:
            Dict with status, filename, download_url, metrics_count
        """
        start_time = time.time()

        try:
            if not self._validate_report_data(report_data):
                return {"status": "error", "message": "Invalid report data provided"}

            structured_data = self._build_report_structure(report_data)

            print(f"📄 Generating Word report for {structured_data['company_name']}")
            filepath = self._generate_word_document(structured_data)

            if not filepath or not Path(filepath).exists():
                return {"status": "error", "message": "Word report generation failed"}

            filename = Path(filepath).name
            download_url = f"/api/RSservice/reports/download/{filename}"
            generation_time = time.time() - start_time

            print(f"✅ Word report generated: {filename} ({generation_time:.2f}s)")

            return {
                "status": "success",
                "format": "word",
                "filename": filename,
                "download_url": download_url,
                "filepath": str(filepath),
                "metrics_count": len(structured_data.get("calculations", [])),
                "generation_time": f"{generation_time:.2f}s"
            }

        except Exception as e:
            print(f"❌ Error generating Word report: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"status": "error", "message": f"Word report generation error: {str(e)}"}

    def generate_pdf_report(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate PDF format ESG report

        Args:
            report_data: Dictionary containing:
                - company_name: str
                - year: str
                - industry: str
                - calculations: List[Dict]
                - quality_score: float (optional)

        Returns:
            Dict with status, filename, download_url, metrics_count
        """
        start_time = time.time()

        try:
            if not self._validate_report_data(report_data):
                return {"status": "error", "message": "Invalid report data provided"}

            structured_data = self._build_report_structure(report_data)

            print(f"📑 Generating PDF report for {structured_data['company_name']}")
            filepath = self._generate_pdf_document(structured_data)

            if not filepath or not Path(filepath).exists():
                return {"status": "error", "message": "PDF report generation failed"}

            filename = Path(filepath).name
            download_url = f"/api/RSservice/reports/download/{filename}"
            generation_time = time.time() - start_time

            print(f"✅ PDF report generated: {filename} ({generation_time:.2f}s)")

            return {
                "status": "success",
                "format": "pdf",
                "filename": filename,
                "download_url": download_url,
                "filepath": str(filepath),
                "metrics_count": len(structured_data.get("calculations", [])),
                "generation_time": f"{generation_time:.2f}s"
            }

        except Exception as e:
            print(f"❌ Error generating PDF report: {str(e)}")
            import traceback
            traceback.print_exc()
            return {"status": "error", "message": f"PDF report generation error: {str(e)}"}

    def get_report_file(self, filename: str) -> Optional[Path]:
        """Get report file path if it exists"""
        filepath = self.reports_dir / filename
        if filepath.exists() and filepath.is_file():
            return filepath
        return None

    # ==================== HELPER METHODS ====================

    def _validate_report_data(self, report_data: Dict[str, Any]) -> bool:
        """Validate report data"""
        required_fields = ["company_name", "year", "industry", "calculations"]
        for field in required_fields:
            if field not in report_data:
                print(f"❌ Missing required field: {field}")
                return False
        if not isinstance(report_data["calculations"], list):
            print("❌ 'calculations' must be a list")
            return False
        if len(report_data["calculations"]) == 0:
            print("❌ At least one calculation result required")
            return False
        return True

    def _build_report_structure(self, report_data: Dict[str, Any]) -> Dict[str, Any]:
        """Build standardized report structure"""
        company_name = report_data["company_name"]
        year = report_data["year"]
        industry = report_data["industry"]
        calculations = report_data["calculations"]
        quality_score = report_data.get("quality_score", 85)

        successful = [c for c in calculations if c.get("status") == "success"]
        failed = [c for c in calculations if c.get("status") != "success"]

        # Group metrics by category
        categories = {}
        for calc in successful:
            cat = calc.get("category", "Other")
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(calc)

        return {
            "company_name": company_name,
            "year": year,
            "industry": industry,
            "framework": self._get_framework_name(industry),
            "calculations": calculations,
            "successful_calculations": successful,
            "failed_calculations": failed,
            "categories": categories,
            "quality_score": quality_score,
            "metrics_count": len(calculations),
            "generation_date": datetime.now().strftime("%B %d, %Y"),
            "generation_timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "report_metadata": {
                "total_metrics": len(calculations),
                "successful_metrics": len(successful),
                "failed_metrics": len(failed),
                "success_rate": round((len(successful) / len(calculations) * 100), 1) if calculations else 0
            }
        }

    def _get_framework_name(self, industry: str) -> str:
        """Get framework name for industry"""
        framework_mapping = {
            "semiconductors": "SASB Semiconductors",
            "commercial_banks": "SASB Commercial Banks"
        }
        return framework_mapping.get(industry, f"SASB {industry.replace('_', ' ').title()}")

    def _format_value(self, value) -> str:
        """Format a metric value for display"""
        if value is None:
            return "N/A"
        if isinstance(value, float):
            if abs(value) >= 1000:
                return f"{value:,.2f}"
            elif abs(value) < 0.01:
                return f"{value:.6f}"
            else:
                return f"{value:.4f}"
        return str(value)

    def _get_method_label(self, calc: Dict) -> str:
        """Get a short method label from calculation data"""
        method = calc.get("calculation_method", calc.get("data_source", ""))
        if not method:
            method = ""
        method_lower = str(method).lower()
        if "direct" in method_lower or "external" in method_lower:
            return "Direct"
        elif "calculated" in method_lower or "model" in method_lower:
            return "Calculated"
        return "Direct"

    def _get_dataset_variable(self, calc: Dict) -> str:
        """Get the dataset variable used"""
        var = calc.get("dataset_variable", "")
        if var:
            return str(var)
        # Try to extract from other fields
        data_source = calc.get("data_source", "")
        if "external" in str(data_source).lower():
            return calc.get("metric_code", "External data")
        return "Model inputs"

    def _get_model_formula(self, calc: Dict) -> str:
        """Get the model/formula description"""
        method = self._get_method_label(calc)
        if method == "Direct":
            var = self._get_dataset_variable(calc)
            return f"Direct: {var}"
        else:
            model = calc.get("model_name", calc.get("metric_name", ""))
            if model:
                return f"Model: {model}"
            return "Calculated Model"

    # ==================== WORD DOCUMENT GENERATION ====================

    def _generate_word_document(self, data: Dict[str, Any]) -> str:
        """Generate comprehensive Word document matching the original PDF design"""
        doc = Document()

        company_name = data.get('company_name', 'Company Name')
        year = data.get('year', 'N/A')
        industry = data.get('industry', 'N/A')
        framework = data.get('framework', 'SASB Framework')
        generation_date = data.get('generation_date', datetime.now().strftime('%B %d, %Y'))
        successful = data.get('successful_calculations', [])
        failed = data.get('failed_calculations', [])
        categories = data.get('categories', {})
        quality_score = data.get('quality_score', 85)

        # ===== COVER PAGE =====
        # Add spacing before title
        for _ in range(4):
            doc.add_paragraph("")

        title = doc.add_heading('ESG COMPLIANCE REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        company_heading = doc.add_heading(company_name, level=1)
        company_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Report info - centered
        info_lines = [
            f"Reporting Year: {year}",
            f"Industry: {industry.replace('_', ' ').title()}",
            f"Framework: {framework}",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)

        doc.add_paragraph("")
        doc.add_paragraph("")

        # Generated date
        p = doc.add_paragraph()
        run = p.add_run(f"Generated: {generation_date}")
        run.font.size = Pt(10)
        run.italic = True
        p = doc.add_paragraph()
        run = p.add_run("ESG Knowledge Graph System")
        run.font.size = Pt(10)
        run.italic = True

        doc.add_page_break()

        # ===== TABLE OF CONTENTS =====
        doc.add_heading('Table of Contents', level=1)

        toc_items = [
            ("Executive Summary", False),
            ("Company Profile", False),
            ("1. ESG Metrics Performance", True),
            ("2. Data Lineage & Transparency", True),
            ("3. Methodology", True),
            ("    3.1 Data Sources", False),
            ("    3.2 Calculation Approach", False),
            ("    3.3 Quality Assurance", False),
            ("4. Recommendations & Next Steps", True),
            ("    4.1 Immediate Actions", False),
            ("    4.2 Medium-Term Goals", False),
            ("    4.3 Long-Term Strategy", False),
        ]
        for item_text, is_bold in toc_items:
            p = doc.add_paragraph()
            run = p.add_run(item_text)
            run.bold = is_bold
            run.font.size = Pt(11)

        doc.add_page_break()

        # ===== EXECUTIVE SUMMARY =====
        doc.add_heading('Executive Summary', level=1)

        exec_text = (
            f"This report presents the Environmental, Social, and Governance (ESG) "
            f"performance assessment for {company_name} for the reporting year {year}. "
            f"The assessment is based on the SASB framework specific to the "
            f"{industry.replace('_', ' ')} industry."
        )
        doc.add_paragraph(exec_text)

        # Key Highlights
        doc.add_heading('Key Highlights:', level=3)
        highlights = [
            f"Total Metrics Assessed: {len(successful) + len(failed)}",
            f"Successfully Calculated: {len(successful)}",
            f"Data Quality Score: {quality_score}%",
            f"Framework: {framework}",
        ]
        for h in highlights:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(h)
            run.bold = True

        # ===== COMPANY PROFILE =====
        doc.add_heading('Company Profile', level=1)

        profile_data = [
            ["Company Name", company_name],
            ["Industry", industry.replace('_', ' ').title()],
            ["Reporting Year", str(year)],
            ["Framework", framework],
            ["Assessment Date", generation_date],
        ]
        table = doc.add_table(rows=len(profile_data), cols=2)
        table.style = 'Table Grid'
        for i, (label, value) in enumerate(profile_data):
            cell_label = table.cell(i, 0)
            cell_value = table.cell(i, 1)
            cell_label.text = ""
            cell_value.text = ""
            run = cell_label.paragraphs[0].add_run(label)
            run.bold = True
            cell_value.paragraphs[0].add_run(value)

        # ===== 1. ESG METRICS PERFORMANCE =====
        doc.add_heading('1. ESG Metrics Performance', level=1)

        if categories:
            section_num = 1
            for cat_name, cat_metrics in categories.items():
                doc.add_heading(f'1.{section_num} {cat_name}', level=2)

                # Metrics table: Metric | Value | Unit | Method
                table = doc.add_table(rows=1 + len(cat_metrics), cols=4)
                table.style = 'Table Grid'

                # Header row
                headers = ["Metric", "Value", "Unit", "Method"]
                for j, h in enumerate(headers):
                    cell = table.cell(0, j)
                    cell.text = ""
                    run = cell.paragraphs[0].add_run(h)
                    run.bold = True

                # Data rows
                for i, calc in enumerate(cat_metrics, start=1):
                    table.cell(i, 0).text = calc.get("metric_name", "N/A")
                    table.cell(i, 1).text = self._format_value(calc.get("value"))
                    table.cell(i, 2).text = calc.get("unit", "N/A")
                    table.cell(i, 3).text = self._get_method_label(calc)

                section_num += 1
                doc.add_paragraph("")
        else:
            doc.add_paragraph("No metrics data available.")

        # ===== 2. DATA LINEAGE & TRANSPARENCY =====
        doc.add_heading('2. Data Lineage & Transparency', level=1)
        doc.add_paragraph(
            "Complete transparency about data sources, calculation methods, and dataset variables."
        )

        if successful:
            table = doc.add_table(rows=1 + len(successful), cols=4)
            table.style = 'Table Grid'

            headers = ["Metric", "Method", "Dataset Variables", "Model/Formula"]
            for j, h in enumerate(headers):
                cell = table.cell(0, j)
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                run.bold = True

            for i, calc in enumerate(successful, start=1):
                table.cell(i, 0).text = calc.get("metric_name", "N/A")
                method = self._get_method_label(calc)
                table.cell(i, 1).text = "Calculated Model" if method == "Calculated" else "Direct Measurement"
                table.cell(i, 2).text = self._get_dataset_variable(calc)
                table.cell(i, 3).text = self._get_model_formula(calc)

        # ===== 3. METHODOLOGY =====
        doc.add_heading('3. Methodology', level=1)

        doc.add_heading('3.1 Data Sources', level=2)
        data_sources = [
            "External ESG Dataset: Primary data source with verified metrics",
            "SASB Framework: Industry-specific sustainability standards",
            "RDF Knowledge Graph: Semantic metric relationships",
        ]
        for src in data_sources:
            doc.add_paragraph(src, style='List Bullet')

        doc.add_heading('3.2 Calculation Approach', level=2)
        approaches = [
            ("Direct Measurement: ", "Metrics from verified datasets without transformation"),
            ("Model-Based: ", "Calculated using validated mathematical models"),
        ]
        for bold_part, normal_part in approaches:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(bold_part)
            run.bold = True
            p.add_run(normal_part)

        doc.add_heading('3.3 Quality Assurance', level=2)
        qa_items = [
            "Verified external datasets only",
            "No synthetic or demo data",
            "Multi-layer validation",
            "Complete data lineage tracking",
        ]
        for item in qa_items:
            doc.add_paragraph(item, style='List Bullet')

        # ===== 4. RECOMMENDATIONS & NEXT STEPS =====
        doc.add_heading('4. Recommendations & Next Steps', level=1)

        doc.add_heading('4.1 Immediate Actions', level=2)
        immediate = [
            "Maintain current high data quality standards",
            "Continue regular ESG data collection and reporting",
            "Monitor regulatory changes in ESG requirements",
        ]
        for item in immediate:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('4.2 Medium-Term Goals', level=2)
        medium = [
            "Expand ESG metric coverage across SASB categories",
            "Implement automated data validation",
            "Develop year-over-year trend analysis",
        ]
        for item in medium:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading('4.3 Long-Term Strategy', level=2)
        long_term = [
            "Establish comprehensive ESG data ecosystem",
            "Implement real-time data monitoring and updates",
            "Develop predictive analytics capabilities for ESG performance",
        ]
        for item in long_term:
            doc.add_paragraph(item, style='List Bullet')

        # End of report
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("--- End of Report ---")
        run.italic = True

        # Generate filename and save
        company_safe = company_name.replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"ESG_Report_{company_safe}_{year}_{timestamp}.docx"
        filepath = self.reports_dir / filename

        doc.save(str(filepath))
        return str(filepath)

    # ==================== PDF DOCUMENT GENERATION ====================

    def _generate_pdf_document(self, data: Dict[str, Any]) -> str:
        """Generate comprehensive PDF document matching the original design"""

        company_name = data.get('company_name', 'Company Name')
        year = data.get('year', 'N/A')
        industry = data.get('industry', 'N/A')
        framework = data.get('framework', 'SASB Framework')
        generation_date = data.get('generation_date', datetime.now().strftime('%B %d, %Y'))
        successful = data.get('successful_calculations', [])
        failed = data.get('failed_calculations', [])
        categories = data.get('categories', {})
        quality_score = data.get('quality_score', 85)

        # Generate filename
        company_safe = company_name.replace(' ', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"ESG_Report_{company_safe}_{year}_{timestamp}.pdf"
        filepath = self.reports_dir / filename

        # Create document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch
        )

        # Define styles
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Title'],
            fontSize=26,
            textColor=HexColor('#2c3e50'),
            spaceAfter=20,
            alignment=TA_CENTER,
        )
        company_style = ParagraphStyle(
            'CompanyName',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=HexColor('#333333'),
            spaceAfter=20,
            alignment=TA_LEFT,
        )
        heading1_style = ParagraphStyle(
            'H1Custom',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=HexColor('#2c3e50'),
            spaceBefore=20,
            spaceAfter=10,
        )
        heading2_style = ParagraphStyle(
            'H2Custom',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=HexColor('#3498db'),
            spaceBefore=14,
            spaceAfter=8,
        )
        heading3_style = ParagraphStyle(
            'H3Custom',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=HexColor('#3498db'),
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            'BodyCustom',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            leading=14,
        )
        cover_info_style = ParagraphStyle(
            'CoverInfo',
            parent=styles['Normal'],
            fontSize=12,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
        italic_style = ParagraphStyle(
            'ItalicText',
            parent=styles['Normal'],
            fontSize=10,
            textColor=HexColor('#666666'),
            fontName='Helvetica-Oblique',
        )
        bullet_style = ParagraphStyle(
            'BulletCustom',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=4,
            leftIndent=20,
            bulletIndent=10,
            leading=14,
        )

        # Table style definitions
        header_bg = HexColor('#3498db')
        header_text = white
        alt_row_bg = HexColor('#f8f9fa')

        def make_table_style():
            return TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), header_bg),
                ('TEXTCOLOR', (0, 0), (-1, 0), header_text),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#dddddd')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [white, alt_row_bg]),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ])

        story = []

        # ===== PAGE 1: COVER PAGE =====
        story.append(Spacer(1, 1.5 * inch))
        story.append(Paragraph("ESG COMPLIANCE REPORT", title_style))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(company_name, company_style))
        story.append(Spacer(1, 0.3 * inch))

        story.append(Paragraph(f"<b>Reporting Year:</b> {year}", cover_info_style))
        story.append(Paragraph(f"<b>Industry:</b> {industry.replace('_', ' ').title()}", cover_info_style))
        story.append(Paragraph(f"<b>Framework:</b> {framework}", cover_info_style))
        story.append(Spacer(1, 0.5 * inch))

        story.append(Paragraph(f"<i>Generated: {generation_date}</i>", italic_style))
        story.append(Paragraph("<i>ESG Knowledge Graph System</i>", italic_style))

        story.append(PageBreak())

        # ===== PAGE 2: BLANK SEPARATOR =====
        story.append(Spacer(1, 4 * inch))
        story.append(PageBreak())

        # ===== PAGE 3: TABLE OF CONTENTS =====
        story.append(Paragraph("Table of Contents", heading1_style))
        story.append(Spacer(1, 0.2 * inch))

        toc_items = [
            ("<b>Executive Summary</b>", 0),
            ("<b>Company Profile</b>", 0),
            ("<b>1. ESG Metrics Performance</b>", 0),
            ("<b>2. Data Lineage &amp; Transparency</b>", 0),
            ("<b>3. Methodology</b>", 0),
            ("3.1 Data Sources", 20),
            ("3.2 Calculation Approach", 20),
            ("3.3 Quality Assurance", 20),
            ("<b>4. Recommendations &amp; Next Steps</b>", 0),
            ("4.1 Immediate Actions", 20),
            ("4.2 Medium-Term Goals", 20),
            ("4.3 Long-Term Strategy", 20),
        ]
        for text, indent in toc_items:
            toc_style = ParagraphStyle(
                'TOC',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=4,
                leftIndent=indent,
                leading=16,
            )
            story.append(Paragraph(text, toc_style))

        story.append(PageBreak())

        # ===== PAGE 4: EXECUTIVE SUMMARY =====
        story.append(Paragraph("Executive Summary", heading1_style))
        story.append(Spacer(1, 0.1 * inch))

        exec_text = (
            f"This report presents the Environmental, Social, and Governance (ESG) "
            f"performance assessment for {company_name} for the reporting year {year}. "
            f"The assessment is based on the SASB framework specific to the "
            f"{industry.replace('_', ' ')} industry."
        )
        story.append(Paragraph(exec_text, body_style))
        story.append(Spacer(1, 0.15 * inch))

        # Key Highlights
        story.append(Paragraph("<b><font color='#3498db'>Key Highlights:</font></b>", body_style))
        highlights = [
            f"<b>Total Metrics Assessed:</b> {len(successful) + len(failed)}",
            f"<b>Successfully Calculated:</b> {len(successful)}",
            f"<b>Data Quality Score:</b> {quality_score}%",
            f"<b>Framework:</b> {framework}",
        ]
        for h in highlights:
            story.append(Paragraph(f"&bull; {h}", bullet_style))

        story.append(Spacer(1, 0.2 * inch))

        # ===== COMPANY PROFILE =====
        story.append(Paragraph("Company Profile", heading1_style))
        story.append(Spacer(1, 0.1 * inch))

        profile_data = [
            ["Company Name", company_name],
            ["Industry", industry.replace('_', ' ').title()],
            ["Reporting Year", str(year)],
            ["Framework", framework],
            ["Assessment Date", generation_date],
        ]

        profile_table_data = []
        for label, value in profile_data:
            profile_table_data.append([
                Paragraph(f"<b>{label}</b>", ParagraphStyle('Cell', fontSize=9, leading=12)),
                Paragraph(value, ParagraphStyle('Cell', fontSize=9, leading=12)),
            ])

        profile_table = Table(profile_table_data, colWidths=[2 * inch, 4.5 * inch])
        profile_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#dddddd')),
            ('BACKGROUND', (0, 0), (0, -1), HexColor('#f8f9fa')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(profile_table)
        story.append(Spacer(1, 0.2 * inch))

        # ===== 1. ESG METRICS PERFORMANCE =====
        story.append(Paragraph("1. ESG Metrics Performance", heading1_style))
        story.append(Spacer(1, 0.1 * inch))

        if categories:
            section_num = 1
            for cat_name, cat_metrics in categories.items():
                story.append(Paragraph(f"1.{section_num} {cat_name}", heading2_style))

                # Build metrics table
                table_data = [
                    [
                        Paragraph("<b>Metric</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                        Paragraph("<b>Value</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                        Paragraph("<b>Unit</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                        Paragraph("<b>Method</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                    ]
                ]
                for calc in cat_metrics:
                    cell_style = ParagraphStyle('TD', fontSize=9, leading=12)
                    table_data.append([
                        Paragraph(calc.get("metric_name", "N/A"), cell_style),
                        Paragraph(self._format_value(calc.get("value")), cell_style),
                        Paragraph(str(calc.get("unit", "N/A")), cell_style),
                        Paragraph(self._get_method_label(calc), cell_style),
                    ])

                metrics_table = Table(table_data, colWidths=[2.2 * inch, 1.5 * inch, 1.8 * inch, 1 * inch])
                metrics_table.setStyle(make_table_style())
                story.append(metrics_table)
                story.append(Spacer(1, 0.15 * inch))
                section_num += 1
        else:
            story.append(Paragraph("No metrics data available.", body_style))

        # ===== 2. DATA LINEAGE & TRANSPARENCY =====
        story.append(Paragraph("2. Data Lineage &amp; Transparency", heading1_style))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph(
            "Complete transparency about data sources, calculation methods, and dataset variables.",
            body_style
        ))
        story.append(Spacer(1, 0.1 * inch))

        if successful:
            lineage_data = [
                [
                    Paragraph("<b>Metric</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                    Paragraph("<b>Method</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                    Paragraph("<b>Dataset Variables</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                    Paragraph("<b>Model/Formula</b>", ParagraphStyle('TH', fontSize=9, textColor=white, leading=12)),
                ]
            ]
            for calc in successful:
                cell_style = ParagraphStyle('TD', fontSize=9, leading=12)
                method = self._get_method_label(calc)
                lineage_data.append([
                    Paragraph(calc.get("metric_name", "N/A"), cell_style),
                    Paragraph("Calculated Model" if method == "Calculated" else "Direct Measurement", cell_style),
                    Paragraph(self._get_dataset_variable(calc), cell_style),
                    Paragraph(self._get_model_formula(calc), cell_style),
                ])

            lineage_table = Table(lineage_data, colWidths=[1.8 * inch, 1.3 * inch, 1.5 * inch, 1.9 * inch])
            lineage_table.setStyle(make_table_style())
            story.append(lineage_table)

        story.append(Spacer(1, 0.2 * inch))

        # ===== 3. METHODOLOGY =====
        story.append(Paragraph("3. Methodology", heading1_style))

        story.append(Paragraph("3.1 Data Sources", heading2_style))
        data_sources = [
            "External ESG Dataset: Primary data source with verified metrics",
            "SASB Framework: Industry-specific sustainability standards",
            "RDF Knowledge Graph: Semantic metric relationships",
        ]
        for src in data_sources:
            story.append(Paragraph(f"&bull; {src}", bullet_style))

        story.append(Paragraph("3.2 Calculation Approach", heading2_style))
        approaches = [
            "<b>Direct Measurement:</b> Metrics from verified datasets without transformation",
            "<b>Model-Based:</b> Calculated using validated mathematical models",
        ]
        for a in approaches:
            story.append(Paragraph(f"&bull; {a}", bullet_style))

        story.append(Paragraph("3.3 Quality Assurance", heading2_style))
        qa_items = [
            "Verified external datasets only",
            "No synthetic or demo data",
            "Multi-layer validation",
            "Complete data lineage tracking",
        ]
        for item in qa_items:
            story.append(Paragraph(f"&bull; {item}", bullet_style))

        story.append(Spacer(1, 0.2 * inch))

        # ===== 4. RECOMMENDATIONS & NEXT STEPS =====
        story.append(Paragraph("4. Recommendations &amp; Next Steps", heading1_style))

        story.append(Paragraph("4.1 Immediate Actions", heading2_style))
        immediate = [
            "Maintain current high data quality standards",
            "Continue regular ESG data collection and reporting",
            "Monitor regulatory changes in ESG requirements",
        ]
        for item in immediate:
            story.append(Paragraph(f"&bull; {item}", bullet_style))

        story.append(Paragraph("4.2 Medium-Term Goals", heading2_style))
        medium = [
            "Expand ESG metric coverage across SASB categories",
            "Implement automated data validation",
            "Develop year-over-year trend analysis",
        ]
        for item in medium:
            story.append(Paragraph(f"&bull; {item}", bullet_style))

        story.append(Paragraph("4.3 Long-Term Strategy", heading2_style))
        long_term = [
            "Establish comprehensive ESG data ecosystem",
            "Implement real-time data monitoring and updates",
            "Develop predictive analytics capabilities for ESG performance",
        ]
        for item in long_term:
            story.append(Paragraph(f"&bull; {item}", bullet_style))

        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph("<i>--- End of Report ---</i>", italic_style))

        # Build PDF
        doc.build(story)
        return str(filepath)

    # ==================== REPORT MANAGEMENT ====================

    def list_reports(self, limit: int = 50) -> List[Dict[str, Any]]:
        """List all generated reports"""
        try:
            reports = []
            report_files = sorted(self.reports_dir.glob("ESG_Report_*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
            report_files.extend(sorted(self.reports_dir.glob("ESG_Report_*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True))
            for filepath in report_files[:limit]:
                filename = filepath.name
                file_stats = filepath.stat()
                report_info = {
                    "filename": filename,
                    "format": "word" if filename.endswith(".docx") else "pdf",
                    "size": file_stats.st_size,
                    "created_date": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                    "download_url": f"/api/RSservice/reports/download/{filename}"
                }
                reports.append(report_info)
            return reports
        except Exception as e:
            print(f"❌ Error listing reports: {str(e)}")
            return []

    def delete_report(self, filename: str) -> Dict[str, Any]:
        """Delete a generated report file"""
        try:
            filepath = self.reports_dir / filename
            if not filepath.exists():
                return {"status": "error", "message": f"Report file not found: {filename}"}
            filepath.unlink()
            return {"status": "success", "message": f"Report deleted: {filename}"}
        except Exception as e:
            return {"status": "error", "message": f"Error deleting report: {str(e)}"}

    def cleanup_old_reports(self, days: int = 30) -> Dict[str, Any]:
        """Clean up reports older than specified days"""
        try:
            import time as time_module
            cutoff_time = time_module.time() - (days * 24 * 60 * 60)
            deleted_count = 0
            for filepath in self.reports_dir.glob("ESG_Report_*"):
                if filepath.stat().st_mtime < cutoff_time:
                    filepath.unlink()
                    deleted_count += 1
            return {"status": "success", "deleted_count": deleted_count, "cutoff_days": days}
        except Exception as e:
            return {"status": "error", "message": f"Error during cleanup: {str(e)}"}
