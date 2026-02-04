#!/usr/bin/env python3
"""
Generate Chapter 4 and Chapter 5 Word documents for PhD thesis.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_chapter4():
    """Generate Chapter 4: ESG Metric Knowledge Graph"""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Chapter Title
    title = doc.add_heading('Chapter 4', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_title = doc.add_heading('ESG Metric Knowledge Graph', level=1)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4.1 Introduction
    doc.add_heading('4.1 Introduction', level=1)

    doc.add_paragraph(
        'Environmental, Social, and Governance (ESG) reporting has become a critical component '
        'of corporate disclosure, driven by increasing regulatory requirements and stakeholder demand '
        'for transparency. Organisations must navigate complex relationships between reporting frameworks, '
        'disclosure categories, metrics, computational models, and underlying data sources. Regulatory '
        'frameworks such as the IFRS Sustainability Disclosure Standards (IFRS S1 and S2), the Task Force '
        'on Climate-related Financial Disclosures (TCFD), and the Sustainability Accounting Standards Board '
        '(SASB) define what organisations should disclose. However, these frameworks often leave open how '
        'metrics are operationalised, computed, and traced back to their data origins. This gap introduces '
        'challenges in metric selection, consistency, transparency, and auditability—particularly when '
        'organisations must comply with multiple frameworks simultaneously.'
    )

    doc.add_paragraph(
        'Knowledge graphs have emerged as a powerful paradigm for representing complex, interconnected '
        'domain knowledge in a machine-interpretable form. By encoding entities and their relationships '
        'as structured semantic networks, knowledge graphs enable systematic querying, reasoning, and '
        'integration across heterogeneous information sources. These capabilities are well-suited to the '
        'challenges of ESG reporting, where traceability, transparency, and interoperability are paramount.'
    )

    doc.add_paragraph(
        'To address these challenges, this chapter introduces the ESG Metric Knowledge Graph (ESGMKG), '
        'a domain-specific knowledge graph designed to explicitly model the semantic relationships among '
        'industries, ESG reporting frameworks, categories, metrics, computational models, implementations, '
        'dataset variables, and data sources. ESGMKG provides a structured representation of ESG reporting '
        'knowledge that enables systematic navigation from high-level industry requirements down to the '
        'concrete data variables and their sources.'
    )

    doc.add_paragraph('ESGMKG serves three key purposes:')

    doc.add_paragraph(
        '1. Semantic Traceability: Establishing explicit links between reported metrics and their '
        'underlying data sources and computational models.', style='List Number'
    )
    doc.add_paragraph(
        '2. Flexible Recomposition: Enabling metric selection and recomposition across different '
        'reporting frameworks.', style='List Number'
    )
    doc.add_paragraph(
        '3. Formal Foundation: Providing a machine-interpretable foundation for querying, reasoning, '
        'and automation in ESG reporting systems.', style='List Number'
    )

    doc.add_paragraph(
        'The remainder of this chapter is organised as follows. Section 4.2 describes the ontology design '
        'methodology, including the hybrid development approach and validation strategy. Section 4.3 presents '
        'the conceptual design of ESGMKG, detailing the core entities, relationships, and modelling decisions. '
        'Section 4.4 demonstrates ESGMKG through instantiation, competency question evaluation, and example '
        'queries. Section 4.5 concludes with a discussion of limitations and future research directions.'
    )

    # 4.2 Design Methodology
    doc.add_heading('4.2 Design Methodology', level=1)

    doc.add_paragraph(
        'This section describes the methodology used to develop ESGMKG, including the ontology engineering '
        'approach and the validation strategy based on competency questions.'
    )

    # 4.2.1 Ontology Engineering Approach
    doc.add_heading('4.2.1 Ontology Engineering Approach', level=2)

    doc.add_paragraph(
        'The development of ESGMKG follows established ontology engineering principles. Specifically, '
        'the ontology adopts a hybrid approach that integrates both top-down and bottom-up design strategies. '
        'This hybrid methodology is necessary because ESG reporting encompasses two distinct types of knowledge: '
        '(1) the stable conceptual structures prescribed by regulatory frameworks, and (2) the heterogeneous '
        'operational practices used by organisations and data providers.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Top-Down Design. ')
    run.bold = True
    p.add_run(
        'The top-down perspective begins with the formal structures defined by ESG reporting frameworks. '
        'Frameworks such as IFRS S1 and TCFD prescribe high-level disclosure requirements organised into '
        'categories such as environmental risk, environmental opportunity, social risk, and governance risk. '
        'These framework-defined structures serve as conceptual anchors that ensure alignment between ESGMKG '
        'and the terminology and logic of regulatory requirements. By deriving the upper levels of the ontology '
        'from authoritative framework definitions, ESGMKG maintains consistency with established ESG reporting standards.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Bottom-Up Design. ')
    run.bold = True
    p.add_run(
        'The bottom-up perspective incorporates the operational realities of ESG data workflows. In practice, '
        'ESG metrics are often the product of multi-step transformations involving estimation, aggregation, or '
        'normalisation. To accurately reflect these complexities, ESGMKG models metrics alongside their '
        'computational models, implementations, dataset variables, and data sources. This bottom-up grounding '
        'ensures that the ontology captures not only what should be reported, but also how reported values are '
        'derived from underlying data.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Hybrid Integration. ')
    run.bold = True
    p.add_run(
        'By combining top-down and bottom-up perspectives, ESGMKG bridges regulatory intent and operational '
        'reality. The resulting ontology captures both the conceptual foundations of ESG reporting and the '
        'practical workflows through which ESG information is produced, validated, and consumed. This hybrid '
        'approach aligns with established recommendations for domain ontology development.'
    )

    # 4.2.2 Validation via Competency Questions
    doc.add_heading('4.2.2 Validation via Competency Questions', level=2)

    doc.add_paragraph(
        'Ontology validation is performed using competency questions (CQs), a widely adopted technique for '
        'verifying that an ontology captures the knowledge required to support its intended use cases. '
        'Competency questions represent realistic information needs that the ontology must be able to answer. '
        'If the instantiated knowledge graph can answer all competency questions consistently and unambiguously, '
        'the ontology is considered validated for its intended purpose.'
    )

    doc.add_paragraph(
        'For ESGMKG, seven competency questions are defined spanning the full knowledge graph structure: '
        'identifying which reporting frameworks apply to a given industry; retrieving the categories included '
        'within a framework; discovering which metrics belong to each category; determining how a metric is '
        'calculated or measured; identifying which metrics serve as inputs for a model; finding which '
        'implementation executes a model; and tracing dataset variables to their original data sources. '
        'Table 4.6 in Section 4.4.2 presents the complete set of competency questions used to validate ESGMKG.'
    )

    # 4.3 Conceptual Design of ESGMKG
    doc.add_heading('4.3 Conceptual Design of ESGMKG', level=1)

    doc.add_paragraph(
        'This section presents the conceptual design of ESGMKG, including the core entities, relationships, '
        'and the rationale behind key modelling decisions.'
    )

    # 4.3.1 Core Entities
    doc.add_heading('4.3.1 Core Entities', level=2)

    doc.add_paragraph(
        'ESGMKG comprises eight core entities that together represent the full chain of information '
        'underpinning ESG reporting. Figure 4.1 illustrates these entities and their relationships.'
    )

    # Placeholder for Figure 4.1
    doc.add_paragraph('[FIGURE 4.1: Conceptual structure of ESGMKG showing eight core entities and their semantic relationships.]')

    p = doc.add_paragraph()
    run = p.add_run('Industry. ')
    run.bold = True
    p.add_run(
        'An Industry represents the sector or domain in which an organisation operates, such as financial '
        'services, energy, or manufacturing. Industries serve as the entry point to the knowledge graph, '
        'as different industries are subject to different reporting frameworks based on regulatory requirements '
        'and stakeholder expectations.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Reporting Framework. ')
    run.bold = True
    p.add_run(
        'A Reporting Framework represents a standardised ESG reporting standard or guideline. Examples include '
        'IFRS S1, TCFD, and SASB. Each framework defines a set of disclosure requirements organised into '
        'reporting categories.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Category. ')
    run.bold = True
    p.add_run(
        'A Category organises ESG disclosures into thematic areas. Common categories include environmental risk, '
        'environmental opportunity, social risk, and governance risk. Categories provide a logical grouping of '
        'related metrics within a framework.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Metric. ')
    run.bold = True
    p.add_run(
        'A Metric represents a measurable quantity used to evaluate ESG performance. Examples include carbon '
        'emissions intensity, water withdrawal, and employee turnover rate. Metrics may be derived through two '
        'distinct paths: they can be calculated by computational models that aggregate or transform input data, '
        'or they can be obtained directly from dataset variables when the metric represents a direct measurement. '
        'Metrics can also serve as inputs to other metrics, enabling the representation of composite indicators. '
        'Section 4.3.5 provides a detailed discussion of the metric schema design.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Model. ')
    run.bold = True
    p.add_run(
        'A Model captures the computational logic used to derive a metric from its inputs. Models specify the '
        'mathematical operations, aggregation rules, or transformation functions applied to input data. By '
        'representing models as explicit entities, ESGMKG ensures that computational processes are transparent '
        'and auditable.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Implementation. ')
    run.bold = True
    p.add_run(
        'An Implementation represents the executable artefact that operationalises a computational model, such '
        'as a Python script, SQL query, or web service. Separating implementations from models allows the same '
        'computational logic to be realised through different technical means, facilitating reuse and platform '
        'independence.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Dataset Variable. ')
    run.bold = True
    p.add_run(
        'A Dataset Variable represents a raw or processed data field used in ESG computations. Examples include '
        'CO2_Scope1, Electricity_Consumed, and Employee_Count. Dataset variables are the atomic data elements '
        'from which metrics are derived.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Data Source. ')
    run.bold = True
    p.add_run(
        'A Data Source identifies where dataset variables originate. Data sources may be internal enterprise '
        'systems (e.g., ERP systems, sustainability management platforms) or external ESG data providers '
        '(e.g., Eurofidai, MSCI, Refinitiv).'
    )

    # 4.3.2 Relationships
    doc.add_heading('4.3.2 Relationships', level=2)

    doc.add_paragraph(
        'The eight core entities are connected through semantically meaningful relationships that enable '
        'navigation and traceability across the knowledge graph. Table 4.1 summarises these relationships.'
    )

    # Table 4.1
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Relationship'
    hdr_cells[1].text = 'Domain'
    hdr_cells[2].text = 'Range'
    hdr_cells[3].text = 'Description'

    # Data rows
    relationships = [
        ('ReportsUsing', 'Industry', 'Framework', 'Industry reports using framework'),
        ('Includes', 'Framework', 'Category', 'Framework includes categories'),
        ('ConsistsOf', 'Category', 'Metric', 'Category consists of metrics'),
        ('IsCalculatedBy', 'Metric', 'Model', 'Metric is calculated by model'),
        ('RequiresInputFrom', 'Model', 'Metric / Variable', 'Model requires inputs'),
        ('ExecutesWith', 'Model', 'Implementation', 'Model executes with implementation'),
        ('ObtainedFrom', 'Metric', 'Variable', 'Metric obtained from variable'),
        ('SourcesFrom', 'Variable', 'Data Source', 'Variable sources from data source'),
    ]

    for i, (rel, domain, range_, desc) in enumerate(relationships, 1):
        row = table.rows[i].cells
        row[0].text = rel
        row[1].text = domain
        row[2].text = range_
        row[3].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.1: ESGMKG relationships connecting core entities.')

    doc.add_paragraph(
        'These relationships collectively enable end-to-end traceability from industry-level reporting '
        'requirements down to the specific data fields that underpin reported outcomes. Notably, metrics '
        'support two derivation paths: the IsCalculatedBy relationship links metrics to computational models '
        'for complex derived indicators, while the ObtainedFrom relationship links metrics directly to dataset '
        'variables for simple measured values. This dual-path design reflects the reality that ESG reporting '
        'includes both computed aggregations and direct measurements.'
    )

    # 4.3.3 Schema Semantics and Traceability
    doc.add_heading('4.3.3 Schema Semantics and Traceability', level=2)

    doc.add_paragraph(
        'The ESGMKG schema is designed around the principle of semantic traceability. Every relationship forms '
        'part of a navigable chain that allows users to trace information from high-level reporting requirements '
        'down to raw data fields. This capability is essential for transparency and auditability in ESG reporting.'
    )

    doc.add_paragraph(
        'The semantics of each relationship are carefully chosen to reflect the underlying domain meaning. '
        'For example, the relationship IsCalculatedBy between a metric and its model reinforces that metrics '
        'are derived through computation rather than directly observed. The relationship ObtainedFrom captures '
        'the alternative case where metrics represent direct measurements from dataset variables. The relationship '
        'RequiresInputFrom captures computational dependencies, while SourcesFrom emphasises data provenance. '
        'By grounding each relationship in a clear semantic role, ESGMKG ensures consistent interpretation '
        'across the ESG reporting lifecycle.'
    )

    # 4.3.4 Design Rationale
    doc.add_heading('4.3.4 Design Rationale', level=2)

    doc.add_paragraph('Several key modelling decisions guide the design of ESGMKG:')

    p = doc.add_paragraph()
    run = p.add_run('Separation of Metrics and Dataset Variables. ')
    run.bold = True
    p.add_run(
        'Metrics are conceptually distinct from dataset variables. Metrics represent the quantities that '
        'organisations report, while dataset variables represent the underlying data fields from which metrics '
        'are derived. This separation allows metric definitions to remain stable even as data sources or '
        'computational methodologies evolve.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Dual Derivation Paths for Metrics. ')
    run.bold = True
    p.add_run(
        'Metrics can be derived through two alternative paths: calculation via models (IsCalculatedBy) or '
        'direct obtainment from dataset variables (ObtainedFrom). Both paths represent valid ways to derive '
        'a metric at the same conceptual level. This design reflects the reality that some ESG metrics require '
        'complex computation (e.g., carbon intensity per revenue) while others represent direct measurements '
        '(e.g., total water withdrawal).'
    )

    p = doc.add_paragraph()
    run = p.add_run('Data Domain Classification. ')
    run.bold = True
    p.add_run(
        'The hasDataDomain property distinguishes between ESG and Financial data sources. This classification '
        'is essential for intensity metrics that combine numerators from ESG sources (e.g., emissions) with '
        'denominators from financial sources (e.g., revenue). By explicitly declaring data domains, the ontology '
        'ensures correct attribution and enables validation of cross-domain calculations.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Explicit Representation of Models. ')
    run.bold = True
    p.add_run(
        'Computational models are represented as first-class entities rather than embedded in metric definitions '
        'or external code. This design improves transparency by making computational logic explicit, enables '
        'model reuse across different metrics, and supports auditability by documenting how reported values '
        'are derived.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Separation of Models and Implementations. ')
    run.bold = True
    p.add_run(
        'Models capture abstract computational logic, while implementations represent concrete executable '
        'artefacts. This separation follows the principle of separating concerns and allows the same model '
        'to be realised through different technical implementations across different platforms or programming '
        'languages.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Recursive Metric Composition. ')
    run.bold = True
    p.add_run(
        'The RequiresInputFrom relationship allows models to take other metrics as inputs, enabling the '
        'representation of composite indicators. This recursive structure is essential for ESG reporting, '
        'where high-level metrics (e.g., overall environmental risk score) often depend on lower-level metrics '
        '(e.g., carbon emissions, water usage).'
    )

    p = doc.add_paragraph()
    run = p.add_run('Framework-Agnostic Design. ')
    run.bold = True
    p.add_run(
        'ESGMKG is designed to be framework-agnostic, capable of representing metrics from IFRS, TCFD, SASB, '
        'GRI, or emerging frameworks without structural modification. This flexibility ensures that the knowledge '
        'graph can accommodate the evolving ESG regulatory landscape.'
    )

    # 4.3.5 Entity Schema Design
    doc.add_heading('4.3.5 Entity Schema Design', level=2)

    doc.add_paragraph(
        'This section presents the detailed schema design for each of the eight ESGMKG entities, specifying '
        'the key properties that characterise each entity type. Understanding these schemas is essential for '
        'implementing and extending the knowledge graph.'
    )

    doc.add_heading('4.3.5.1 Industry and Framework Entities', level=3)

    doc.add_paragraph(
        'The Industry entity represents a business sector classification. Table 4.7a presents the Industry schema.'
    )

    # Table 4.7a: Industry Schema
    table_industry = doc.add_table(rows=4, cols=3)
    table_industry.style = 'Table Grid'

    hdr = table_industry.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    industry_props = [
        ('rdfs:label', 'String', 'Industry name (e.g., "Semiconductors", "Commercial Banks")'),
        ('esg:industryCode', 'String', 'Standard industry classification code'),
        ('esg:reportsUsing', 'URI (Framework)', 'Links to applicable reporting framework'),
    ]

    for i, (prop, dtype, desc) in enumerate(industry_props, 1):
        row = table_industry.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7a: Industry entity schema.')

    doc.add_paragraph(
        'The ReportingFramework entity represents an ESG disclosure standard. Table 4.7b presents the schema.'
    )

    # Table 4.7b: Framework Schema
    table_framework = doc.add_table(rows=5, cols=3)
    table_framework.style = 'Table Grid'

    hdr = table_framework.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    framework_props = [
        ('rdfs:label', 'String', 'Framework name (e.g., "SASB Semiconductors")'),
        ('esg:sourceDocument', 'String', 'Reference to official standard document'),
        ('esg:frameworkVersion', 'String', 'Version identifier of the framework'),
        ('esg:includes', 'URI (Category)', 'Links to disclosure categories within the framework'),
    ]

    for i, (prop, dtype, desc) in enumerate(framework_props, 1):
        row = table_framework.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7b: ReportingFramework entity schema.')

    doc.add_heading('4.3.5.2 Category Entity', level=3)

    doc.add_paragraph(
        'The Category entity represents a disclosure topic grouping within a framework. Table 4.7c presents the schema.'
    )

    # Table 4.7c: Category Schema
    table_category = doc.add_table(rows=4, cols=3)
    table_category.style = 'Table Grid'

    hdr = table_category.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    category_props = [
        ('rdfs:label', 'String', 'Category name (e.g., "Greenhouse Gas Emissions")'),
        ('esg:categoryCode', 'String', 'Standard category identifier'),
        ('esg:consistsOf', 'URI (Metric)', 'Links to metrics within this category'),
    ]

    for i, (prop, dtype, desc) in enumerate(category_props, 1):
        row = table_category.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7c: Category entity schema.')

    doc.add_heading('4.3.5.3 Model and Implementation Entities', level=3)

    doc.add_paragraph(
        'The Model entity captures computational logic for derived metrics. Table 4.7d presents the schema.'
    )

    # Table 4.7d: Model Schema
    table_model = doc.add_table(rows=6, cols=3)
    table_model.style = 'Table Grid'

    hdr = table_model.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    model_props = [
        ('rdfs:label', 'String', 'Model name (e.g., "GHGEmissionIntensityModel")'),
        ('esg:formula', 'String', 'Mathematical formula (e.g., "(scope1 + scope2) / revenue")'),
        ('esg:calculationType', 'String', 'Type of calculation (e.g., "intensity_ratio", "percentage")'),
        ('esg:requiresInputFrom', 'URI (Metric/Variable)', 'Links to required inputs'),
        ('esg:executesWith', 'URI (Implementation)', 'Links to executable implementation'),
    ]

    for i, (prop, dtype, desc) in enumerate(model_props, 1):
        row = table_model.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7d: Model entity schema.')

    doc.add_paragraph(
        'The Implementation entity represents executable code. Table 4.7e presents the schema.'
    )

    # Table 4.7e: Implementation Schema
    table_impl = doc.add_table(rows=5, cols=3)
    table_impl.style = 'Table Grid'

    hdr = table_impl.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    impl_props = [
        ('rdfs:label', 'String', 'Implementation name'),
        ('esg:filePath', 'String', 'Path to executable file'),
        ('esg:functionName', 'String', 'Name of callable function'),
        ('esg:language', 'String', 'Programming language (e.g., "Python")'),
    ]

    for i, (prop, dtype, desc) in enumerate(impl_props, 1):
        row = table_impl.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7e: Implementation entity schema.')

    doc.add_heading('4.3.5.4 DatasetVariable and DataSource Entities', level=3)

    doc.add_paragraph(
        'The DatasetVariable entity represents a raw data field. Table 4.7f presents the schema.'
    )

    # Table 4.7f: DatasetVariable Schema
    table_variable = doc.add_table(rows=6, cols=3)
    table_variable.style = 'Table Grid'

    hdr = table_variable.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    variable_props = [
        ('rdfs:label', 'String', 'Variable label (e.g., "CO2 Direct Scope 1")'),
        ('esg:variableName', 'String', 'Technical variable name (e.g., "CO2DIRECTSCOPE1")'),
        ('esg:alignmentReason', 'String', 'Explanation of how variable maps to SASB requirements'),
        ('esg:confidenceScore', 'Decimal', 'Mapping reliability score (0-100)'),
        ('esg:sourcesFrom', 'URI (DataSource)', 'Links to originating data source'),
    ]

    for i, (prop, dtype, desc) in enumerate(variable_props, 1):
        row = table_variable.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7f: DatasetVariable entity schema.')

    doc.add_paragraph(
        'The DataSource entity identifies data origins. Table 4.7g presents the schema.'
    )

    # Table 4.7g: DataSource Schema
    table_datasource = doc.add_table(rows=5, cols=3)
    table_datasource.style = 'Table Grid'

    hdr = table_datasource.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Data Type'
    hdr[2].text = 'Description'

    datasource_props = [
        ('rdfs:label', 'String', 'Data source name (e.g., "Eurofidai", "WRDS")'),
        ('esg:provider', 'String', 'Organisation providing the data'),
        ('esg:dataFormat', 'String', 'Format of source data (e.g., "CSV")'),
        ('esg:accessMethod', 'String', 'How data is accessed (e.g., "file", "API")'),
    ]

    for i, (prop, dtype, desc) in enumerate(datasource_props, 1):
        row = table_datasource.rows[i].cells
        row[0].text = prop
        row[1].text = dtype
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.7g: DataSource entity schema.')

    # 4.3.6 Metric Schema Design (previously 4.3.5)
    doc.add_heading('4.3.6 Metric Schema Design', level=2)

    doc.add_paragraph(
        'The Metric entity is central to ESGMKG and requires careful schema design to accurately represent '
        'the diverse ways ESG metrics are defined, derived, and sourced. This section presents the metric '
        'schema design, clarifying the classification approach and the key properties that characterise each metric.'
    )

    doc.add_heading('4.3.6.1 Metric Classification by Derivation Method', level=3)

    doc.add_paragraph(
        'A fundamental design decision in ESGMKG is the classification of metrics by their derivation method. '
        'Rather than creating separate entity types for different kinds of metrics, ESGMKG uses a single Metric '
        'entity with a hasCalculationMethod property that indicates how the metric value is obtained. This approach '
        'provides several advantages: it maintains a unified metric representation, enables consistent querying '
        'across all metrics, and allows the same metric to potentially support multiple derivation methods in '
        'different contexts.'
    )

    doc.add_paragraph('The hasCalculationMethod property accepts two values:')

    p = doc.add_paragraph()
    run = p.add_run('direct_measurement: ')
    run.bold = True
    p.add_run(
        'The metric value is obtained directly from a dataset variable without requiring computation. '
        'Direct measurement metrics use the ObtainedFrom relationship to link to their source DatasetVariable. '
        'Examples include total water withdrawal, total energy consumed, and gross Scope 1 emissions, where '
        'the reported value comes directly from measured or reported data.'
    )

    p = doc.add_paragraph()
    run = p.add_run('calculation_model: ')
    run.bold = True
    p.add_run(
        'The metric value is derived through a computational model that transforms or aggregates input data. '
        'Calculated metrics use the IsCalculatedBy relationship to link to their Model entity. Examples include '
        'GHG emission intensity (emissions divided by revenue), percentage of renewable energy (renewable energy '
        'divided by total energy), and water stress ratios.'
    )

    doc.add_paragraph(
        'Figure 4.2 illustrates both derivation paths through a comprehensive example. The left side demonstrates '
        'the calculated metric path with recursive resolution: an aggregate metric (EnvironmentalRiskMetric) depends '
        'on component metrics (GHGEmissionIntensity and AirQualityPollutant), which themselves require inputs from '
        'metrics across ESG and Financial domains, ultimately traced to dataset variables and data sources. The right '
        'side shows the direct metric path: TotalWaterConsumed is obtained directly from a dataset variable without '
        'requiring any computational model. Both paths represent valid approaches—they are alternative ways to derive '
        'a metric value, not a hierarchy of metric types.'
    )

    doc.add_paragraph('[FIGURE 4.2: Metric derivation paths showing calculated metric path with recursive resolution (left) and direct metric path (right).]')

    doc.add_heading('4.3.6.2 Data Domain Classification', level=3)

    doc.add_paragraph(
        'ESG intensity metrics often require data from multiple domains. For example, GHG emission intensity '
        'is calculated as total emissions (an ESG metric) divided by revenue (a financial metric). To support '
        'such cross-domain calculations while maintaining clear data provenance, ESGMKG introduces the '
        'hasDataDomain property.'
    )

    doc.add_paragraph('The hasDataDomain property accepts two values:')

    p = doc.add_paragraph()
    run = p.add_run('ESG: ')
    run.bold = True
    p.add_run(
        'The metric is derived from environmental, social, or governance data sources. ESG domain metrics '
        'include emissions data, water consumption, energy usage, waste generation, and similar sustainability '
        'measurements. These metrics are typically sourced from ESG data providers such as Eurofidai, MSCI, '
        'or corporate sustainability reports.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Financial: ')
    run.bold = True
    p.add_run(
        'The metric is derived from financial data sources. Financial domain metrics include revenue, total '
        'assets, market capitalisation, and similar financial measurements. These metrics are typically sourced '
        'from financial data providers such as WRDS, Bloomberg, or corporate financial statements.'
    )

    doc.add_paragraph(
        'The hasDataDomain property serves several purposes:'
    )

    doc.add_paragraph(
        'Data Provenance: Clearly identifies the source domain for each metric, enabling audit trails that '
        'distinguish between ESG and financial data origins.', style='List Bullet'
    )
    doc.add_paragraph(
        'Data Integration: Guides the system in routing data requests to appropriate data sources '
        '(ESG repositories vs. financial repositories).', style='List Bullet'
    )
    doc.add_paragraph(
        'Validation: Enables validation rules that ensure intensity calculations correctly combine ESG '
        'numerators with financial denominators.', style='List Bullet'
    )

    doc.add_heading('4.3.6.3 Metric Schema Summary', level=3)

    doc.add_paragraph(
        'Table 4.8 summarises the key properties of the Metric entity in ESGMKG.'
    )

    # Table 4.8: Metric Properties
    table_metric = doc.add_table(rows=7, cols=3)
    table_metric.style = 'Table Grid'

    hdr = table_metric.rows[0].cells
    hdr[0].text = 'Property'
    hdr[1].text = 'Values'
    hdr[2].text = 'Description'

    metric_props = [
        ('hasCalculationMethod', 'direct_measurement | calculation_model', 'Indicates how the metric value is derived'),
        ('hasDataDomain', 'ESG | Financial', 'Indicates the source domain of the underlying data'),
        ('hasDescription', 'String', 'Human-readable description of the metric'),
        ('hasUnit', 'String', 'Unit of measurement (e.g., "tonnes CO2-e", "Percentage")'),
        ('hasMetricType', 'Quantitative | Discussion', 'Indicates whether metric is numeric or qualitative'),
        ('hasType', 'SASBRequirement | Manual', 'Indicates whether metric is defined by SASB or custom'),
    ]

    for i, (prop, values, desc) in enumerate(metric_props, 1):
        row = table_metric.rows[i].cells
        row[0].text = prop
        row[1].text = values
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 4.8: Key properties of the Metric entity.')

    doc.add_paragraph(
        'This schema design ensures that all metrics—whether directly measured or calculated, whether from '
        'ESG or financial sources—are represented consistently within a single unified entity type. The '
        'hasCalculationMethod property determines the derivation path, while hasDataDomain enables proper '
        'data source routing and provenance tracking. This approach avoids the complexity of multiple metric '
        'subclasses while providing the semantic richness needed for automated ESG reporting workflows.'
    )

    # 4.4 ESGMKG Instantiation and Evaluation
    doc.add_heading('4.4 ESGMKG Instantiation and Evaluation', level=1)

    doc.add_paragraph(
        'This section demonstrates ESGMKG through instantiation with a representative example, evaluates the '
        'ontology against competency questions, and illustrates querying capabilities using SPARQL.'
    )

    # 4.4.1 Example Instantiation
    doc.add_heading('4.4.1 Example Instantiation', level=2)

    doc.add_paragraph(
        'To demonstrate ESGMKG in practice, we instantiate the knowledge graph with a realistic ESG reporting '
        'scenario based on the SASB framework for the semiconductor industry. The example focuses on environmental '
        'metrics, illustrating both calculated metrics derived through computational models and direct metrics '
        'obtained directly from dataset variables.'
    )

    # 4.4.1.1 Metric Derivation Paths
    doc.add_heading('4.4.1.1 Metric Derivation Paths', level=3)

    doc.add_paragraph('ESGMKG supports two alternative derivation paths for metrics, both at the same conceptual level:')

    doc.add_paragraph(
        'Calculated Metric Path: A metric is derived through a computational model using the IsCalculatedBy '
        'relationship. The model may require other metrics or dataset variables as inputs via the RequiresInputFrom '
        'relationship. When inputs are themselves calculated metrics, the system recursively resolves each input '
        'until all values are obtained from dataset variables.', style='List Bullet'
    )
    doc.add_paragraph(
        'Direct Metric Path: A metric is obtained directly from a dataset variable using the ObtainedFrom '
        'relationship, without requiring any computational model.', style='List Bullet'
    )

    doc.add_paragraph(
        'As illustrated in Figure 4.2, the left side shows a calculated metric with recursive '
        'resolution, where the target metric (EnvironmentalRiskMetric) depends on component metrics that themselves '
        'require computation before resolving to dataset variables. The right side shows the direct metric path '
        'where a metric is obtained directly from a dataset variable without intermediate computation.'
    )

    # 4.4.1.2 Calculated Metric Path
    doc.add_heading('4.4.1.2 Calculated Metric Path: Recursive Resolution', level=3)

    doc.add_paragraph(
        'The calculated metric path in Figure 4.2 demonstrates recursive resolution through an aggregate '
        'environmental risk metric. The resolution proceeds through the following steps:'
    )

    doc.add_paragraph(
        'Step 1 - Target Metric: EnvironmentalRiskMetric is the target metric to be calculated. It uses '
        'EnvironmentalRiskModel which requires two component metrics as inputs.', style='List Bullet'
    )
    doc.add_paragraph(
        'Step 2 - Component Metrics: GHGEmissionIntensity and AirQualityPollutant are themselves calculated '
        'metrics. Each requires its own model and inputs, triggering recursive resolution.', style='List Bullet'
    )
    doc.add_paragraph(
        'Step 3 - Metric Inputs: The component metrics require inputs from multiple domains. GHGEmissionIntensity '
        'requires Scope1, Scope2 (ESG domain) and Revenue (Financial domain). AirQualityPollutant requires '
        'SOX, NOX, and VOC emissions data (ESG domain).', style='List Bullet'
    )
    doc.add_paragraph(
        'Step 4 - Dataset Variables: Each input is mapped to dataset variables (CO2SCOPE1, CO2SCOPE2, revt, '
        'SOXEMISSIONS, NOXEMISSIONS, VOCEMISSIONS) via the ObtainedFrom relationship.', style='List Bullet'
    )
    doc.add_paragraph(
        'Step 5 - Data Sources: Variables are traced to their originating data sources (Eurofidai for ESG '
        'data, WRDS for Financial data) via the SourcesFrom relationship.', style='List Bullet'
    )

    # 4.4.1.3 Direct Metric Path
    doc.add_heading('4.4.1.3 Direct Metric Path', level=3)

    doc.add_paragraph(
        'The direct metric path demonstrates the alternative derivation approach: TotalWaterConsumed is obtained '
        'directly from the WATERCONSUMPTIONTOTAL dataset variable via the ObtainedFrom relationship. No '
        'computational model is required for this path.'
    )

    # 4.4.1.4 Instantiation Tables
    doc.add_heading('4.4.1.4 Instantiation Tables', level=3)

    doc.add_paragraph('The following tables present the complete ESGMKG instantiation.')

    doc.add_paragraph('Table 4.2 shows the reporting path from industry through framework and category to metrics.')

    # Table 4.2
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'

    hdr = table2.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Source Entity'
    hdr[2].text = 'Relationship'
    hdr[3].text = 'Target Entity'

    data2 = [
        ('1', 'Industry: Semiconductors', 'ReportsUsing', 'Framework: SASB Semiconductors'),
        ('2', 'Framework: SASB Semiconductors', 'Includes', 'Category: Environmental Risk'),
        ('3', 'Category: Environmental Risk', 'ConsistsOf', 'Metric: EnvironmentalRiskMetric'),
        ('4', 'Category: Water Management', 'ConsistsOf', 'Metric: TotalWaterConsumed'),
    ]

    for i, (step, src, rel, tgt) in enumerate(data2, 1):
        row = table2.rows[i].cells
        row[0].text = step
        row[1].text = src
        row[2].text = rel
        row[3].text = tgt

    doc.add_paragraph()
    doc.add_paragraph('Table 4.2: ESGMKG instantiation: Reporting path from industry to metrics.')

    doc.add_paragraph('Table 4.3 shows the calculated metric path with recursive resolution, corresponding to '
                      'the left side of Figure 4.2. The table demonstrates both branches: GHGEmissionIntensity '
                      '(requiring Scope1, Scope2, Revenue) and AirQualityPollutant (requiring SOX, NOX, VOC).')

    # Table 4.3
    table3 = doc.add_table(rows=19, cols=4)
    table3.style = 'Table Grid'

    hdr = table3.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Source Entity'
    hdr[2].text = 'Relationship'
    hdr[3].text = 'Target Entity'

    data3 = [
        # Step 1: Target metric and model
        ('5', 'Metric: EnvironmentalRiskMetric', 'IsCalculatedBy', 'Model: EnvironmentalRiskModel'),
        ('6', 'Model: EnvironmentalRiskModel', 'RequiresInputFrom', 'Metric: GHGEmissionIntensity'),
        ('7', 'Model: EnvironmentalRiskModel', 'RequiresInputFrom', 'Metric: AirQualityPollutant'),
        # Step 2: GHG branch resolution
        ('8', 'Metric: GHGEmissionIntensity', 'IsCalculatedBy', 'Model: GHGIntensityModel'),
        ('9', 'Model: GHGIntensityModel', 'RequiresInputFrom', 'Metric: Scope1 (ESG)'),
        ('10', 'Model: GHGIntensityModel', 'RequiresInputFrom', 'Metric: Scope2 (ESG)'),
        ('11', 'Model: GHGIntensityModel', 'RequiresInputFrom', 'Metric: Revenue (Financial)'),
        # Step 2: AirQuality branch resolution
        ('12', 'Metric: AirQualityPollutant', 'IsCalculatedBy', 'Model: AirQualityModel'),
        ('13', 'Model: AirQualityModel', 'RequiresInputFrom', 'Metric: SOX (ESG)'),
        ('14', 'Model: AirQualityModel', 'RequiresInputFrom', 'Metric: NOX (ESG)'),
        ('15', 'Model: AirQualityModel', 'RequiresInputFrom', 'Metric: VOC (ESG)'),
        # Step 3-4: Map metrics to dataset variables
        ('16', 'Metric: Scope1', 'ObtainedFrom', 'Variable: CO2DIRECTSCOPE1'),
        ('17', 'Metric: Scope2', 'ObtainedFrom', 'Variable: CO2INDIRECTSCOPE2'),
        ('18', 'Metric: Revenue', 'ObtainedFrom', 'Variable: revt'),
        ('19', 'Metric: SOX', 'ObtainedFrom', 'Variable: SOXEMISSIONS'),
        ('20', 'Metric: NOX', 'ObtainedFrom', 'Variable: NOXEMISSIONS'),
        ('21', 'Metric: VOC', 'ObtainedFrom', 'Variable: VOCEMISSIONS'),
    ]

    for i, (step, src, rel, tgt) in enumerate(data3, 1):
        row = table3.rows[i].cells
        row[0].text = step
        row[1].text = src
        row[2].text = rel
        row[3].text = tgt

    doc.add_paragraph()
    doc.add_paragraph('Table 4.3: ESGMKG instantiation: Calculated metric path with recursive resolution (both GHG and Air Quality branches).')

    doc.add_paragraph('Table 4.4 shows the direct metric path.')

    # Table 4.4
    table4 = doc.add_table(rows=2, cols=4)
    table4.style = 'Table Grid'

    hdr = table4.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Source Entity'
    hdr[2].text = 'Relationship'
    hdr[3].text = 'Target Entity'

    row = table4.rows[1].cells
    row[0].text = '22'
    row[1].text = 'Metric: TotalWaterConsumed'
    row[2].text = 'ObtainedFrom'
    row[3].text = 'Variable: WATERCONSUMPTIONTOTAL'

    doc.add_paragraph()
    doc.add_paragraph('Table 4.4: ESGMKG instantiation: Direct metric path (no computational model required).')

    doc.add_paragraph('Table 4.5 shows the data provenance path, tracing variables to their data sources.')

    # Table 4.5
    table5 = doc.add_table(rows=9, cols=4)
    table5.style = 'Table Grid'

    hdr = table5.rows[0].cells
    hdr[0].text = 'Step'
    hdr[1].text = 'Source Entity'
    hdr[2].text = 'Relationship'
    hdr[3].text = 'Target Entity'

    data5 = [
        ('23', 'Variable: CO2DIRECTSCOPE1', 'SourcesFrom', 'DataSource: Eurofidai'),
        ('24', 'Variable: CO2INDIRECTSCOPE2', 'SourcesFrom', 'DataSource: Eurofidai'),
        ('25', 'Variable: revt', 'SourcesFrom', 'DataSource: WRDS'),
        ('26', 'Variable: SOXEMISSIONS', 'SourcesFrom', 'DataSource: Eurofidai'),
        ('27', 'Variable: NOXEMISSIONS', 'SourcesFrom', 'DataSource: Eurofidai'),
        ('28', 'Variable: VOCEMISSIONS', 'SourcesFrom', 'DataSource: Eurofidai'),
        ('29', 'Variable: WATERCONSUMPTIONTOTAL', 'SourcesFrom', 'DataSource: Eurofidai'),
        ('30', 'Model: EnvironmentalRiskModel', 'ExecutesWith', 'Impl: environmental_risk_model.py'),
    ]

    for i, (step, src, rel, tgt) in enumerate(data5, 1):
        row = table5.rows[i].cells
        row[0].text = step
        row[1].text = src
        row[2].text = rel
        row[3].text = tgt

    doc.add_paragraph()
    doc.add_paragraph('Table 4.5: ESGMKG instantiation: Data provenance linking variables to data sources.')

    # 4.4.1.5 Summary
    doc.add_heading('4.4.1.5 Summary of Instantiation', level=3)

    doc.add_paragraph('The instantiation demonstrates how ESGMKG captures the full complexity of ESG metric derivation:')

    doc.add_paragraph(
        'Dual derivation paths: Calculated metrics use IsCalculatedBy while direct metrics use ObtainedFrom, '
        'representing alternative approaches at the same conceptual level.', style='List Bullet'
    )
    doc.add_paragraph(
        'Recursive resolution: EnvironmentalRiskMetric demonstrates recursive resolution where calculated metrics '
        '(GHGEmissionIntensity, AirQualityPollutant) depend on other metrics that must be resolved first.', style='List Bullet'
    )
    doc.add_paragraph(
        'Cross-domain composition: GHGEmissionIntensity combines ESG metrics (Scope1, Scope2) with Financial metrics '
        '(Revenue), with hasDataDomain clearly distinguishing source domains.', style='List Bullet'
    )
    doc.add_paragraph(
        'Data provenance: All variables trace back to their data sources (Eurofidai for ESG, WRDS for Financial).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Complete coverage: All 8 entity types and all 8 relationships are demonstrated.', style='List Bullet'
    )

    # 4.4.2 Competency Question Evaluation
    doc.add_heading('4.4.2 Competency Question Evaluation', level=2)

    doc.add_paragraph(
        'The instantiated ESGMKG is validated using competency questions that test its ability to support '
        'essential ESG reporting queries. Table 4.6 presents the seven competency questions and their expected outputs.'
    )

    # Table 4.6
    table6 = doc.add_table(rows=8, cols=3)
    table6.style = 'Table Grid'

    hdr = table6.rows[0].cells
    hdr[0].text = 'CQ'
    hdr[1].text = 'Question'
    hdr[2].text = 'Expected Output'

    cqs = [
        ('CQ1', 'Which framework applies to a given industry?', 'Reporting Framework'),
        ('CQ2', 'What categories are included within that framework?', 'Category'),
        ('CQ3', 'Which metrics belong to each category?', 'Metric'),
        ('CQ4', 'How is a metric calculated or measured?', 'Model or Dataset Variable'),
        ('CQ5', 'Which metrics are inputs for a model?', 'Metric'),
        ('CQ6', 'Which implementation executes a model?', 'Implementation'),
        ('CQ7', 'What is the original data source for a variable?', 'Data Source'),
    ]

    for i, (cq, question, output) in enumerate(cqs, 1):
        row = table6.rows[i].cells
        row[0].text = cq
        row[1].text = question
        row[2].text = output

    doc.add_paragraph()
    doc.add_paragraph('Table 4.6: Competency questions for ESGMKG validation.')

    doc.add_paragraph(
        'All seven competency questions can be answered by the instantiated knowledge graph. For CQ1, querying '
        'which framework applies to the semiconductor industry returns SASB Semiconductors (Table 4.2, Step 1). '
        'For CQ3, querying which metrics belong to the Environmental Risk category returns EnvironmentalRiskMetric '
        'and related metrics. For CQ4, querying how EnvironmentalRiskMetric is derived returns EnvironmentalRiskModel '
        'via IsCalculatedBy (Table 4.3, Step 5), while querying how TotalWaterConsumed is derived returns the '
        'WATERCONSUMPTIONTOTAL dataset variable via ObtainedFrom (Table 4.4, Step 16)—demonstrating both derivation paths. '
        'For CQ5, querying which metrics are inputs for EnvironmentalRiskModel returns GHGEmissionIntensity and '
        'AirQualityPollutant (Table 4.3, Steps 6-7), demonstrating the recursive composition capability.'
    )

    # 4.4.3 Usage Scenarios
    doc.add_heading('4.4.3 Usage Scenarios', level=2)

    doc.add_paragraph(
        'ESGMKG supports diverse usage scenarios relevant to ESG reporting practice. In framework selection, '
        'an organisation identifies applicable reporting frameworks by querying ESGMKG with their industry; '
        'for example, a semiconductor company can discover that SASB Semiconductors is the relevant framework '
        'for their sector. In metric discovery, a sustainability officer explores reporting categories to '
        'identify available metrics and understand their derivation paths, revealing not only what metrics '
        'exist but also whether they are calculated or directly obtained. In audit trail scenarios, an auditor '
        'traces a reported metric back through its derivation path to the underlying dataset variables and data '
        'sources, supporting assurance and regulatory compliance. In model reuse scenarios, a data engineer '
        'identifies computational models that can be reused across different metrics or frameworks, along with '
        'their existing implementations, reducing development effort and ensuring consistency.'
    )

    # 4.4.4 RDF Representation and SPARQL Querying
    doc.add_heading('4.4.4 RDF Representation and SPARQL Querying', level=2)

    doc.add_paragraph(
        'ESGMKG is implemented using the Resource Description Framework (RDF) and the Web Ontology Language (OWL), '
        'ensuring interoperability with semantic web technologies. Each entity and relationship is expressed as '
        'RDF triples that can be queried using SPARQL. Table 4.11 illustrates example RDF triples representing '
        'key entities and relationships from the instantiation.'
    )

    # Table 4.7
    table7 = doc.add_table(rows=10, cols=3)
    table7.style = 'Table Grid'

    hdr = table7.rows[0].cells
    hdr[0].text = 'Subject'
    hdr[1].text = 'Predicate'
    hdr[2].text = 'Object'

    triples = [
        ('esg:semiconductors', 'rdf:type', 'esg:Industry'),
        ('esg:semiconductors', 'esg:ReportsUsing', 'esg:SASBSemiconductors'),
        ('esg:SASBSemiconductors', 'esg:Includes', 'esg:GreenhouseGasEmissions'),
        ('esg:GreenhouseGasEmissions', 'esg:ConsistsOf', 'esg:GHGEmissionIntensity'),
        ('esg:GHGEmissionIntensity', 'esg:IsCalculatedBy', 'esg:GHGEmissionIntensityModel'),
        ('esg:GHGEmissionIntensityModel', 'esg:RequiresInputFrom', 'esg:GrossGlobalScope1Emissions'),
        ('esg:GrossGlobalScope1Emissions', 'esg:hasDataDomain', '"ESG"'),
        ('esg:Revenue', 'esg:hasDataDomain', '"Financial"'),
        ('esg:CO2DIRECTSCOPE1', 'esg:SourcesFrom', 'esg:Eurofidai'),
    ]

    for i, (subj, pred, obj) in enumerate(triples, 1):
        row = table7.rows[i].cells
        row[0].text = subj
        row[1].text = pred
        row[2].text = obj

    doc.add_paragraph()
    doc.add_paragraph('Table 4.11: Example RDF triples in ESGMKG.')

    doc.add_paragraph(
        'SPARQL queries enable flexible retrieval of information from ESGMKG. Table 4.12 shows a query that '
        'retrieves all metrics in a given category along with their derivation method, demonstrating how to '
        'distinguish between calculated and direct metrics.'
    )

    # Table 4.9: Algorithm 1 - SPARQL Query
    table_alg1 = doc.add_table(rows=10, cols=2)
    table_alg1.style = 'Table Grid'

    alg1_rows = [
        ('Line', 'SPARQL Query'),
        ('1', 'PREFIX esg: <http://example.org/esg#>'),
        ('2', ''),
        ('3', 'SELECT ?metric ?model ?variable'),
        ('4', 'WHERE {'),
        ('5', '    esg:EnvironmentalRisk esg:ConsistsOf ?metric .'),
        ('6', '    OPTIONAL { ?metric esg:IsCalculatedBy ?model . }'),
        ('7', '    OPTIONAL { ?metric esg:ObtainedFrom ?variable . }'),
        ('8', '}'),
        ('', 'Note: Line 6 matches calculated metrics; Line 7 matches direct metrics'),
    ]

    for i, (line, code) in enumerate(alg1_rows):
        row = table_alg1.rows[i].cells
        row[0].text = line
        row[1].text = code

    doc.add_paragraph()
    doc.add_paragraph('Table 4.12: SPARQL query to retrieve metrics and their derivation methods (Algorithm 1).')

    doc.add_paragraph(
        'Table 4.13 presents a query that traces a calculated metric through its recursive composition to the '
        'underlying data sources, resolving component metrics along the way.'
    )

    # Table 4.10: Algorithm 2 - SPARQL Query
    table_alg2 = doc.add_table(rows=12, cols=2)
    table_alg2.style = 'Table Grid'

    alg2_rows = [
        ('Line', 'SPARQL Query'),
        ('1', 'PREFIX esg: <http://example.org/esg#>'),
        ('2', ''),
        ('3', 'SELECT ?topMetric ?componentMetric ?variable ?dataSource'),
        ('4', 'WHERE {'),
        ('5', '    ?topMetric esg:IsCalculatedBy ?topModel .          # Target metric'),
        ('6', '    ?topModel esg:RequiresInputFrom ?componentMetric .'),
        ('7', '    ?componentMetric esg:IsCalculatedBy ?compModel .   # Component metric'),
        ('8', '    ?compModel esg:RequiresInputFrom ?variable .'),
        ('9', '    ?variable esg:SourcesFrom ?dataSource .            # Data source'),
        ('10', '}'),
        ('', 'Note: Traces from target metric through component metrics to data sources'),
    ]

    for i, (line, code) in enumerate(alg2_rows):
        row = table_alg2.rows[i].cells
        row[0].text = line
        row[1].text = code

    doc.add_paragraph()
    doc.add_paragraph('Table 4.13: SPARQL query to trace metric hierarchy to data sources (Algorithm 2).')

    doc.add_paragraph(
        'These queries illustrate how ESGMKG enables systematic navigation through both metric derivation '
        'paths, from high-level metrics down to underlying data sources.'
    )

    # 4.5 Discussion and Conclusion
    doc.add_heading('4.5 Discussion and Conclusion', level=1)

    # 4.5.1 Summary of Contributions
    doc.add_heading('4.5.1 Summary of Contributions', level=2)

    doc.add_paragraph(
        'This chapter presented ESGMKG, a domain-specific knowledge graph for ESG metric management. The primary '
        'contribution is a conceptual model comprising eight core entities—Industry, Reporting Framework, Category, '
        'Metric, Model, Implementation, Dataset Variable, and Data Source—connected through eight semantic '
        'relationships that enable end-to-end traceability. A key design feature is the dual derivation paths for '
        'metrics: calculated metrics are linked to computational models via IsCalculatedBy, while direct metrics '
        'are linked to dataset variables via ObtainedFrom. Both paths represent alternative approaches at the same '
        'conceptual level. The hasDataDomain property enables clear distinction between ESG and Financial data '
        'sources, which is essential for intensity metrics that combine data from multiple domains. The hierarchical '
        'metric composition capability, demonstrated within the calculated metric path, shows how aggregate metrics '
        'can depend on component metrics, which in turn depend on raw dataset variables. The ontology was developed '
        'using a hybrid design methodology that combines top-down structures derived from regulatory frameworks with '
        'bottom-up grounding in operational ESG data workflows. The chapter validated the ontology through seven '
        'competency questions representing realistic ESG reporting information needs.'
    )

    # 4.5.2 Limitations
    doc.add_heading('4.5.2 Limitations', level=2)

    doc.add_paragraph(
        'Several limitations of the current work should be acknowledged. The scope of instantiation focuses on '
        'environmental risk metrics for demonstration purposes; comprehensive deployment would require extensive '
        'population with metrics from multiple categories (social, governance) and frameworks across different '
        'industries. The schema does not explicitly model temporal dynamics such as framework versioning, metric '
        'evolution, or historical changes to computational models; incorporating temporal semantics would enable '
        'tracking of how ESG reporting requirements evolve over time. While ESGMKG traces metrics to data sources, '
        'it does not capture data quality metadata such as completeness, accuracy, or timeliness; extending the '
        'ontology with data quality dimensions would enhance its utility for assurance purposes. Finally, the '
        'current approach relies on manual population of the knowledge graph; future work could explore automated '
        'extraction of ESG knowledge from regulatory documents and corporate reports using natural language '
        'processing techniques.'
    )

    # 4.5.3 Future Research Directions
    doc.add_heading('4.5.3 Future Research Directions', level=2)

    doc.add_paragraph(
        'Future research directions include extended framework coverage to incorporate additional frameworks such '
        'as GRI, CDP, and TNFD, along with social and governance metrics; integration with large language models '
        'to enable natural language querying and automated knowledge extraction from unstructured ESG documents; '
        'reasoning and inference capabilities using OWL reasoning to infer implicit relationships, detect '
        'inconsistencies, and recommend appropriate metrics based on industry context; and enterprise integration '
        'patterns for connecting ESGMKG with enterprise sustainability management systems and ESG reporting platforms.'
    )

    # 4.5.4 Conclusion
    doc.add_heading('4.5.4 Conclusion', level=2)

    doc.add_paragraph(
        'ESGMKG addresses the critical need for structured, machine-interpretable representations of ESG reporting '
        'knowledge. By explicitly modelling the relationships between frameworks, metrics, models, and data '
        'sources—and by supporting both calculated and direct metric derivation paths with hierarchical '
        'composition—ESGMKG enables transparency, traceability, and interoperability in ESG reporting. The '
        'hasDataDomain property ensures that metrics from heterogeneous sources (ESG and Financial) are correctly '
        'attributed, enabling accurate intensity calculations. The knowledge graph provides a foundation for '
        'building intelligent ESG reporting systems that can navigate the complexity of multi-framework compliance '
        'while maintaining auditability and consistency.'
    )

    # Save document
    doc.save('/Users/mingqin/Downloads/esg-knowledge-graph-demo/Demo corresponding acadedmic paper/Chapter4_ESG_Metric_Knowledge_Graph.docx')
    print("Chapter 4 saved successfully!")


def create_chapter5():
    """Generate Chapter 5: ESG Reporting Framework Implementation"""
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Chapter Title
    title = doc.add_heading('Chapter 5', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    main_title = doc.add_heading('ESG Reporting Framework Implementation', level=1)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5.1 Introduction
    doc.add_heading('5.1 Introduction', level=1)

    doc.add_paragraph(
        'Building upon the ESGMKG ontology presented in Chapter 4, this chapter details the implementation '
        'of an ESG reporting framework that operationalises the conceptual model through a service-oriented '
        'architecture. The framework demonstrates how knowledge graph-based approaches can address practical '
        'challenges in ESG metric calculation, data integration, and report generation. By bridging ontological '
        'design with software engineering practice, the implementation validates the conceptual model while '
        'providing a functional prototype for ESG reporting automation. As outlined in the research methodology '
        '(Chapter 3, Section 3.4), this work constitutes the second contribution of this thesis. Whereas '
        'Chapter 4 addressed the first contribution—the design and formalisation of the ESGMKG ontology—the '
        'present chapter focuses on the practical realisation of that ontology through executable software services.'
    )

    doc.add_paragraph(
        'The framework serves as a proof-of-concept that demonstrates the feasibility and utility of '
        'knowledge graph-based ESG reporting, following the Design Science Research methodology adopted '
        'in this thesis (Chapter 3, Section 3.2). It addresses three key requirements: (1) semantic navigation '
        'through the knowledge graph to discover applicable metrics for a given industry and framework; '
        '(2) automated metric calculation using both direct measurement retrieval and computational model '
        'execution; and (3) comprehensive report generation that provides full traceability from reported '
        'values to underlying data sources.'
    )

    doc.add_paragraph(
        'The remainder of this chapter is organised as follows. Section 5.2 defines the system requirements '
        'and use cases. Section 5.3 describes the overall system architecture. Section 5.4 presents the service '
        'layer design with four microservices. Section 5.5 details the API layer. Section 5.6 describes the data '
        'layer components, including external data input specifications and internal data formats. Section 5.7 '
        'presents the user interface design. Section 5.8 presents sequence diagrams illustrating key use case '
        'workflows. Section 5.9 discusses scalability considerations. Section 5.10 concludes with a summary.'
    )

    # 5.2 Requirements and Use Cases
    doc.add_heading('5.2 Requirements and Use Cases', level=1)

    doc.add_paragraph(
        'This section defines the functional and non-functional requirements for the ESG reporting framework, '
        'along with the primary use cases that guide system design and implementation.'
    )

    # 5.2.1 Functional Requirements
    doc.add_heading('5.2.1 Functional Requirements', level=2)

    doc.add_paragraph(
        'The ESG reporting framework must satisfy the following functional requirements derived from the '
        'ESGMKG competency questions and practical ESG reporting workflows:'
    )

    # Functional Requirements Table
    table_fr = doc.add_table(rows=9, cols=3)
    table_fr.style = 'Table Grid'

    hdr = table_fr.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'
    hdr[2].text = 'Related CQ'

    fr_data = [
        ('FR1', 'The system shall enable users to discover applicable reporting frameworks for a given industry.', 'CQ1'),
        ('FR2', 'The system shall display disclosure categories within a selected reporting framework.', 'CQ2'),
        ('FR3', 'The system shall list metrics belonging to each disclosure category with their properties.', 'CQ3'),
        ('FR4', 'The system shall determine and display the calculation method (direct or calculated) for each metric.', 'CQ4'),
        ('FR5', 'The system shall identify and retrieve input metrics/variables required for calculated metrics.', 'CQ5'),
        ('FR6', 'The system shall execute appropriate implementations to compute calculated metric values.', 'CQ6'),
        ('FR7', 'The system shall trace and display the data provenance from metrics to original data sources.', 'CQ7'),
        ('FR8', 'The system shall generate comprehensive ESG reports with full audit trails.', '-'),
    ]

    for i, (id_, req, cq) in enumerate(fr_data, 1):
        row = table_fr.rows[i].cells
        row[0].text = id_
        row[1].text = req
        row[2].text = cq

    doc.add_paragraph()
    doc.add_paragraph('Table 5.1: Functional requirements aligned with competency questions.')

    # 5.2.2 Non-Functional Requirements
    doc.add_heading('5.2.2 Non-Functional Requirements', level=2)

    doc.add_paragraph('The system shall satisfy the following non-functional requirements:')

    # Non-Functional Requirements Table
    table_nfr = doc.add_table(rows=6, cols=2)
    table_nfr.style = 'Table Grid'

    hdr = table_nfr.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = 'Requirement'

    nfr_data = [
        ('NFR1', 'Extensibility: The system shall support addition of new frameworks, metrics, and data sources without code changes.'),
        ('NFR2', 'Traceability: All calculated values shall be traceable to their underlying data sources and computational models.'),
        ('NFR3', 'Interoperability: The system shall use standard semantic web technologies (RDF, SPARQL, OWL).'),
        ('NFR4', 'Usability: The user interface shall provide step-by-step guidance through the ESG reporting workflow.'),
        ('NFR5', 'Performance: The system shall respond to API requests within 2 seconds under normal operating conditions.'),
    ]

    for i, (id_, req) in enumerate(nfr_data, 1):
        row = table_nfr.rows[i].cells
        row[0].text = id_
        row[1].text = req

    doc.add_paragraph()
    doc.add_paragraph('Table 5.2: Non-functional requirements.')

    # 5.2.3 Use Cases
    doc.add_heading('5.2.3 Use Cases', level=2)

    doc.add_paragraph(
        'The framework supports three core use cases that address the primary workflows in ESG reporting: '
        'report generation, metric calculation, and data access. These use cases reflect core challenges '
        'faced by ESG analysts and map directly to the competency questions defined in Chapter 4.'
    )

    # UC1: ESG Report Generation
    doc.add_heading('5.2.3.1 UC1: ESG Report Generation', level=3)

    doc.add_paragraph(
        'This use case automates the production of ESG reports aligned with relevant reporting frameworks. '
        'The workflow begins when a user selects an industry classification, triggering framework identification '
        'through CQ1. Once a framework is selected, the system enumerates disclosure categories via CQ2 and '
        'identifies required metrics through CQ3. For each metric, the system invokes UC2 (Metric Calculation) '
        'to obtain values, then assembles the complete report with framework-compliant structure and full '
        'data lineage documentation.'
    )

    doc.add_paragraph('Actor: ESG Reporting Officer / Sustainability Analyst')
    doc.add_paragraph('Precondition: User has identified their industry sector and target company.')
    doc.add_paragraph('Main Flow:')
    doc.add_paragraph('1. User selects industry classification (e.g., Semiconductors).', style='List Number')
    doc.add_paragraph('2. System identifies applicable reporting frameworks via CQ1.', style='List Number')
    doc.add_paragraph('3. User selects framework (e.g., SASB Semiconductors).', style='List Number')
    doc.add_paragraph('4. System enumerates disclosure categories via CQ2.', style='List Number')
    doc.add_paragraph('5. For each category, system identifies required metrics via CQ3.', style='List Number')
    doc.add_paragraph('6. System invokes UC2 to calculate each metric value.', style='List Number')
    doc.add_paragraph('7. System assembles report with metric values, methodology, and data lineage.', style='List Number')
    doc.add_paragraph('Postcondition: Framework-compliant ESG report generated with full audit trail.')
    doc.add_paragraph('Related Competency Questions: CQ1, CQ2, CQ3')

    # UC2: Metric Calculation
    doc.add_heading('5.2.3.2 UC2: Metric Calculation', level=3)

    doc.add_paragraph(
        'This use case supports the transparent computation of individual ESG metrics. When a metric '
        'calculation is requested, the system first determines whether the metric value is obtained '
        'directly from a dataset variable (hasCalculationMethod: direct_measurement) or derived through '
        'a computational model (hasCalculationMethod: calculation_model) using CQ4. For calculated metrics, '
        'the system identifies required inputs via CQ5, recursively resolves any dependent metrics, and '
        'executes the appropriate model implementation identified through CQ6. The calculation process '
        'maintains complete provenance, recording how each value was derived.'
    )

    doc.add_paragraph('Actor: ESG Reporting Officer / Data Analyst')
    doc.add_paragraph('Precondition: Metric has been identified for calculation (typically via UC1).')
    doc.add_paragraph('Main Flow:')
    doc.add_paragraph('1. System determines calculation method via CQ4 (direct_measurement or calculation_model).', style='List Number')
    doc.add_paragraph('2. For direct metrics:', style='List Number')
    doc.add_paragraph('   - System identifies source dataset variable via ObtainedFrom relationship', style='List Bullet')
    doc.add_paragraph('   - System invokes UC3 to retrieve value from data source', style='List Bullet')
    doc.add_paragraph('3. For calculated metrics:', style='List Number')
    doc.add_paragraph('   - System identifies required inputs via CQ5 (metrics or variables)', style='List Bullet')
    doc.add_paragraph('   - System recursively calculates any dependent metrics', style='List Bullet')
    doc.add_paragraph('   - System identifies model implementation via CQ6', style='List Bullet')
    doc.add_paragraph('   - System executes calculation with resolved inputs', style='List Bullet')
    doc.add_paragraph('4. System returns calculated value with complete derivation provenance.', style='List Number')
    doc.add_paragraph('Postcondition: Metric value calculated with transparent methodology and provenance.')
    doc.add_paragraph('Related Competency Questions: CQ4, CQ5, CQ6')

    # UC3: Data Access
    doc.add_heading('5.2.3.3 UC3: Data Access', level=3)

    doc.add_paragraph(
        'This use case traces the raw sources of metric inputs through direct access to ESG datasets '
        'and metadata. For each dataset variable involved in a metric, the system invokes CQ7 to '
        'identify the original data source, retrieves the requested value, and records provenance '
        'metadata including source attribution, temporal coverage, and data quality indicators. '
        'This use case ensures that all reported values can be traced back to their authoritative sources.'
    )

    doc.add_paragraph('Actor: Data Analyst / Auditor')
    doc.add_paragraph('Precondition: Dataset variable has been identified (typically via UC2).')
    doc.add_paragraph('Main Flow:')
    doc.add_paragraph('1. System identifies data source for variable via CQ7 (e.g., Eurofidai, WRDS).', style='List Number')
    doc.add_paragraph('2. System routes request based on hasDataDomain (ESG or Financial).', style='List Number')
    doc.add_paragraph('3. System retrieves value for specified company and time period.', style='List Number')
    doc.add_paragraph('4. System records provenance metadata (source, timestamp, quality indicators).', style='List Number')
    doc.add_paragraph('5. System returns value with complete source attribution.', style='List Number')
    doc.add_paragraph('Postcondition: Data value retrieved with full provenance to original source.')
    doc.add_paragraph('Related Competency Questions: CQ7')

    # Use Case to CQ Mapping Table
    doc.add_paragraph()
    doc.add_paragraph('Table 5.3 provides the detailed mapping between use cases and competency questions.')

    table_uc = doc.add_table(rows=4, cols=3)
    table_uc.style = 'Table Grid'

    hdr = table_uc.rows[0].cells
    hdr[0].text = 'Use Case'
    hdr[1].text = 'Competency Questions'
    hdr[2].text = 'Primary Output'

    uc_data = [
        ('UC1: ESG Report Generation', 'CQ1, CQ2, CQ3', 'Framework-compliant ESG report'),
        ('UC2: Metric Calculation', 'CQ4, CQ5, CQ6', 'Calculated metric values with provenance'),
        ('UC3: Data Access', 'CQ7', 'Raw data values with source attribution'),
    ]

    for i, (uc, cqs, output) in enumerate(uc_data, 1):
        row = table_uc.rows[i].cells
        row[0].text = uc
        row[1].text = cqs
        row[2].text = output

    doc.add_paragraph()
    doc.add_paragraph('Table 5.3: Use Case to Competency Question Mapping.')

    # Use Case Diagram placeholder
    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.1: Use case diagram showing the three core ESG reporting use cases (UC1: Report Generation, UC2: Metric Calculation, UC3: Data Access) and their relationships to competency questions CQ1-CQ7.]')

    # 5.3 System Architecture
    doc.add_heading('5.3 System Architecture', level=1)

    doc.add_paragraph(
        'The ESG reporting framework follows a four-layer service-oriented architecture that separates '
        'concerns and enables modularity. Figure 5.2 illustrates the overall system architecture.'
    )

    doc.add_paragraph('[FIGURE 5.2: Four-layer system architecture showing Presentation, API Gateway, Service, and Data layers.]')

    doc.add_heading('5.3.1 Architectural Layers', level=2)

    doc.add_paragraph('The system comprises four distinct layers:')

    p = doc.add_paragraph()
    run = p.add_run('Presentation Layer. ')
    run.bold = True
    p.add_run(
        'A web-based user interface that enables users to select industries, frameworks, and metrics, '
        'then view calculated results and generated reports. The interface provides a step-by-step '
        'workflow guiding users through the ESG reporting process.'
    )

    p = doc.add_paragraph()
    run = p.add_run('API Gateway Layer. ')
    run.bold = True
    p.add_run(
        'A RESTful API implemented using Flask that exposes service functionality to the presentation layer '
        'and external consumers. The API gateway handles request routing, response formatting, and error handling.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Service Layer. ')
    run.bold = True
    p.add_run(
        'Four microservices that implement the core business logic: Knowledge Graph Service for ontology '
        'navigation, Data Retrieval Service for data access, Calculation Service for metric computation, '
        'and Report Service for summary generation.'
    )

    p = doc.add_paragraph()
    run = p.add_run('Data Layer. ')
    run.bold = True
    p.add_run(
        'Persistent storage comprising the RDF knowledge graph (stored in Turtle format), ESG data repositories '
        '(CSV files from Eurofidai), financial data repositories (CSV files from WRDS), and the model repository '
        'containing calculation model definitions and implementations.'
    )

    # 5.4 Service Layer Design
    doc.add_heading('5.4 Service Layer Design', level=1)

    doc.add_paragraph(
        'The service layer implements four microservices, each with a specific responsibility aligned with '
        'the ESGMKG ontology structure. These services interact through well-defined interfaces to support '
        'the ESG reporting workflow.'
    )

    # 5.4.1 Knowledge Graph Service
    doc.add_heading('5.4.1 Knowledge Graph Service', level=2)

    doc.add_paragraph(
        'The Knowledge Graph Service provides semantic navigation capabilities over the ESGMKG ontology. '
        'It loads the RDF knowledge graph from Turtle format and executes SPARQL queries to answer the '
        'competency questions defined in Chapter 4.'
    )

    doc.add_paragraph('The service implements the following query operations aligned with competency questions:')

    # Table: KG Service Operations
    table_kg = doc.add_table(rows=8, cols=3)
    table_kg.style = 'Table Grid'

    hdr = table_kg.rows[0].cells
    hdr[0].text = 'Operation'
    hdr[1].text = 'Competency Question'
    hdr[2].text = 'Description'

    kg_ops = [
        ('get_frameworks_for_industry()', 'CQ1', 'Returns applicable reporting frameworks for an industry'),
        ('get_categories_for_framework()', 'CQ2', 'Returns disclosure categories within a framework'),
        ('get_metrics_for_category()', 'CQ3', 'Returns metrics belonging to a category'),
        ('get_calculation_method()', 'CQ4', 'Returns model or variable for metric derivation'),
        ('get_model_inputs()', 'CQ5', 'Returns input metrics/variables for a model'),
        ('get_model_implementation()', 'CQ6', 'Returns implementation details for a model'),
        ('get_variable_source()', 'CQ7', 'Returns data source for a dataset variable'),
    ]

    for i, (op, cq, desc) in enumerate(kg_ops, 1):
        row = table_kg.rows[i].cells
        row[0].text = op
        row[1].text = cq
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 5.4: Knowledge Graph Service operations aligned with competency questions.')

    doc.add_paragraph(
        'The service uses RDFLib to parse and query the knowledge graph, providing an abstraction layer '
        'that shields other services from SPARQL syntax details.'
    )

    # 5.4.2 Data Retrieval Service
    doc.add_heading('5.4.2 Data Retrieval Service', level=2)

    doc.add_paragraph(
        'The Data Retrieval Service manages access to underlying data repositories, retrieving values for '
        'dataset variables from their respective data sources. The service implements the data provenance '
        'chain: Metric → ObtainedFrom → DatasetVariable → SourcesFrom → DataSource.'
    )

    doc.add_paragraph('The service provides the following capabilities:')

    doc.add_paragraph(
        'Company Discovery: Retrieves available companies by industry sector from the data repositories.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Variable Retrieval: Fetches values for dataset variables given a company and time period.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Multi-Source Integration: Handles data from multiple sources (Eurofidai for ESG, WRDS for Financial) '
        'using the hasDataDomain property to route requests appropriately.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Unit Conversion: Applies necessary unit conversions based on metadata in the knowledge graph.',
        style='List Bullet'
    )

    # 5.4.3 Calculation Service
    doc.add_heading('5.4.3 Calculation Service', level=2)

    doc.add_paragraph(
        'The Calculation Service implements the metric derivation logic defined in ESGMKG. It handles both '
        'direct measurement retrieval and computational model execution, following the dual derivation paths '
        'described in Chapter 4.'
    )

    doc.add_paragraph('The calculation workflow proceeds as follows:')

    doc.add_paragraph(
        'Step 1 (CQ4): Query the knowledge graph to determine the calculation method for the requested metric '
        '(direct_measurement or calculation_model).', style='List Number'
    )
    doc.add_paragraph(
        'Step 2a - Direct Path: If hasCalculationMethod is "direct_measurement", query CQ7 to identify the '
        'data source, then retrieve the value via the Data Retrieval Service.', style='List Number'
    )
    doc.add_paragraph(
        'Step 2b - Calculated Path: If hasCalculationMethod is "calculation_model", proceed to Step 3.',
        style='List Number'
    )
    doc.add_paragraph(
        'Step 3 (CQ5): Query the knowledge graph to identify input metrics/variables required by the model.',
        style='List Number'
    )
    doc.add_paragraph(
        'Step 4: For each input, recursively apply Steps 1-3 to obtain values (handles metrics from both '
        'ESG and Financial domains based on hasDataDomain).', style='List Number'
    )
    doc.add_paragraph(
        'Step 5 (CQ6): Query the knowledge graph to identify the model implementation.',
        style='List Number'
    )
    doc.add_paragraph(
        'Step 6: Execute the implementation with the gathered input values and return the calculated result.',
        style='List Number'
    )

    doc.add_paragraph('Figure 5.4 illustrates the calculation service workflow.')

    doc.add_paragraph('[FIGURE 5.4: Calculation Service workflow showing dual derivation paths and recursive input resolution.]')

    doc.add_paragraph(
        'Error handling is implemented at each step. If any required input is unavailable or a calculation '
        'fails, the service returns a detailed error message indicating the failure point while continuing '
        'to process other metrics where possible.'
    )

    # 5.4.4 Report Service
    doc.add_heading('5.4.4 Report Service', level=2)

    doc.add_paragraph(
        'The Report Service aggregates calculation results into comprehensive ESG reports. It provides '
        'summary statistics, visualisations, and full audit trails documenting how each reported value '
        'was derived.'
    )

    doc.add_paragraph('The service generates reports containing:')

    doc.add_paragraph(
        'Metric Summary: Calculated values for all requested metrics with units and descriptions.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Derivation Audit Trail: For each metric, the complete chain from reported value through model '
        'execution to underlying data sources.', style='List Bullet'
    )
    doc.add_paragraph(
        'Data Provenance: Source attribution for all dataset variables used in calculations.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Visualisations: Charts and graphs summarising ESG performance across categories.',
        style='List Bullet'
    )

    # 5.5 API Layer
    doc.add_heading('5.5 API Layer', level=1)

    doc.add_paragraph(
        'The API layer exposes service functionality through RESTful endpoints. Table 5.5 summarises the '
        'available API endpoints.'
    )

    # Table 5.5: API Endpoints
    table_api = doc.add_table(rows=10, cols=4)
    table_api.style = 'Table Grid'

    hdr = table_api.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Method'
    hdr[2].text = 'Description'
    hdr[3].text = 'Related CQ'

    api_endpoints = [
        ('/industries', 'GET', 'List all available industries', '—'),
        ('/industries/{id}/frameworks', 'GET', 'Retrieve reporting frameworks for an industry', 'CQ1'),
        ('/frameworks/{id}/categories', 'GET', 'Retrieve disclosure categories within a framework', 'CQ2'),
        ('/categories/{id}/metrics', 'GET', 'Retrieve metrics within a category', 'CQ3'),
        ('/metrics/{id}', 'GET', 'Retrieve metric definition and calculation method', 'CQ4'),
        ('/metrics/{id}/inputs', 'GET', 'Retrieve required inputs for a derived metric', 'CQ5'),
        ('/calculate', 'POST', 'Compute metric value for a specified company and period', 'CQ6'),
        ('/calculations/{id}/lineage', 'GET', 'Retrieve provenance and lineage for a calculation', 'CQ7'),
        ('/reports/generate', 'POST', 'Generate an ESG report with aggregated results', 'UC3'),
    ]

    for i, (endpoint, method, desc, cq) in enumerate(api_endpoints, 1):
        row = table_api.rows[i].cells
        row[0].text = endpoint
        row[1].text = method
        row[2].text = desc
        row[3].text = cq

    doc.add_paragraph()
    doc.add_paragraph('Table 5.5: REST API endpoints mapped to competency questions.')

    doc.add_paragraph(
        'The API follows resource-oriented REST conventions, using domain concepts (industries, frameworks, '
        'categories, metrics) as URL nouns rather than exposing internal service structure. This design '
        'decouples clients from implementation details and enables backend refactoring without breaking '
        'API contracts. All responses follow a consistent JSON format including status, data payload, '
        'and error information where applicable.'
    )

    # 5.6 Data Layer
    doc.add_heading('5.6 Data Layer', level=1)

    doc.add_paragraph(
        'The data layer comprises three core components: the knowledge graph repository storing ESGMKG entities '
        'and relationships, the external data sources providing raw ESG and financial data, and the model '
        'repository containing calculation definitions. This section details each component and specifies '
        'the data formats required for system integration.'
    )

    # 5.6.1 Knowledge Graph Repository
    doc.add_heading('5.6.1 Knowledge Graph Repository', level=2)

    doc.add_paragraph(
        'The knowledge graph is stored in RDF Turtle (.ttl) format, implementing the ESGMKG ontology defined in '
        'Chapter 4. The graph is loaded into memory using RDFLib and queried via SPARQL. The current '
        'instantiation includes:'
    )

    doc.add_paragraph('2 Industries: Commercial Banks, Semiconductors', style='List Bullet')
    doc.add_paragraph('2 Reporting Frameworks: SASB Commercial Banks, SASB Semiconductors', style='List Bullet')
    doc.add_paragraph('14 Disclosure Categories spanning environmental, social, and governance topics', style='List Bullet')
    doc.add_paragraph('29 Metrics with hasCalculationMethod and hasDataDomain properties', style='List Bullet')
    doc.add_paragraph('6 Computational Models with mathematical formulas and input specifications', style='List Bullet')
    doc.add_paragraph('6 Implementations with Python function bindings', style='List Bullet')
    doc.add_paragraph('15 Dataset Variables with alignment metadata', style='List Bullet')
    doc.add_paragraph('2 Data Sources: Eurofidai (ESG), WRDS (Financial)', style='List Bullet')

    doc.add_paragraph(
        'The knowledge graph file is located at data/rdf/esg_knowledge_graph.ttl and is automatically loaded '
        'by the UnifiedKnowledgeGraphService at system startup.'
    )

    # 5.6.2 External Data Sources
    doc.add_heading('5.6.2 External Data Sources', level=2)

    doc.add_paragraph(
        'The system integrates data from external ESG and financial data providers. This section documents '
        'the data sources, their CSV formats, and the requirements for adding new data sources to the system.'
    )

    doc.add_heading('5.6.2.1 Available Data Sources', level=3)

    doc.add_paragraph('Table 5.6 summarises the external data sources currently integrated into the system.')

    # Data Sources Table
    table_ds = doc.add_table(rows=4, cols=4)
    table_ds.style = 'Table Grid'

    hdr = table_ds.rows[0].cells
    hdr[0].text = 'Data Source'
    hdr[1].text = 'Data Domain'
    hdr[2].text = 'File Location'
    hdr[3].text = 'Record Count'

    ds_data = [
        ('Eurofidai (Semiconductors)', 'ESG', 'data/External dataset/Semiconductors_Eurofidai_EnvironmentData.csv', '~105,000'),
        ('Eurofidai (Commercial Banks)', 'ESG', 'data/External dataset/Commercial_Banks_Eurofidai_EnvironmentData.csv', '~317,000'),
        ('WRDS Compustat', 'Financial', 'data/External dataset/Semiconductor_WRDS_FinancialData.csv', '~500'),
    ]

    for i, (src, domain, location, count) in enumerate(ds_data, 1):
        row = table_ds.rows[i].cells
        row[0].text = src
        row[1].text = domain
        row[2].text = location
        row[3].text = count

    doc.add_paragraph()
    doc.add_paragraph('Table 5.6: External data sources integrated into the system.')

    doc.add_heading('5.6.2.2 Eurofidai ESG Data Format', level=3)

    doc.add_paragraph(
        'ESG data from Eurofidai (sourced from Clarity AI) uses a "long format" CSV structure where each row '
        'represents a single metric observation for a company. This normalised format enables flexible storage '
        'of diverse metrics. Table 5.7 specifies the CSV schema.'
    )

    # Eurofidai CSV Schema Table
    table_csv_esg = doc.add_table(rows=11, cols=3)
    table_csv_esg.style = 'Table Grid'

    hdr = table_csv_esg.rows[0].cells
    hdr[0].text = 'Column'
    hdr[1].text = 'Description'
    hdr[2].text = 'Example'

    csv_esg_schema = [
        ('company_name', 'Company legal name', 'NVIDIA Corp'),
        ('perm_id', 'Permanent company identifier', '4295914405'),
        ('metric_name', 'Variable identifier for ESGMKG mapping', 'SOXEMISSIONS'),
        ('metric_value', 'Numeric value of the metric', '125.5'),
        ('metric_unit', 'Unit of measurement', 'Tons of SOx'),
        ('metric_year', 'Reporting period end date', '2022-12-31'),
        ('disclosure', 'Data origin (REPORTED/CALCULATED/ESTIMATED)', 'ESTIMATED'),
        ('pillar', 'ESG pillar (E, S, G)', 'E'),
        ('provider_name', 'Data provider', 'Clarity AI'),
        ('industry', 'Industry classification', 'Semiconductors'),
    ]

    for i, (col, desc, example) in enumerate(csv_esg_schema, 1):
        row = table_csv_esg.rows[i].cells
        row[0].text = col
        row[1].text = desc
        row[2].text = example

    doc.add_paragraph()
    doc.add_paragraph('Table 5.7: Eurofidai ESG data CSV schema.')

    doc.add_paragraph('Example rows from the Eurofidai dataset:')
    doc.add_paragraph(
        'company_name,metric_name,metric_value,metric_unit,metric_year,pillar\n'
        'NVIDIA Corp,CO2DIRECTSCOPE1,125000.5,Tons of CO2,2022-12-31,E\n'
        'NVIDIA Corp,SOXEMISSIONS,0.04,Tons of SOx,2022-12-31,E\n'
        'NVIDIA Corp,VOCEMISSIONS,12.5,Tons of VOC,2022-12-31,E',
        style='Normal'
    )

    doc.add_paragraph(
        'Key metric_name values used by ESGMKG: CO2DIRECTSCOPE1, CO2INDIRECTSCOPE2 (GHG emissions), '
        'SOXEMISSIONS, NOXEMISSIONS, VOCEMISSIONS (air quality), WATERCONSUMPTIONTOTAL, ENERGYUSETOTAL, HAZARDOUSWASTE.'
    )

    doc.add_heading('5.6.2.3 WRDS Financial Data Format', level=3)

    doc.add_paragraph(
        'Financial data from WRDS Compustat uses a similar long format. Table 5.8 specifies the schema.'
    )

    # WRDS CSV Schema Table
    table_csv_fin = doc.add_table(rows=6, cols=3)
    table_csv_fin.style = 'Table Grid'

    hdr = table_csv_fin.rows[0].cells
    hdr[0].text = 'Column'
    hdr[1].text = 'Description'
    hdr[2].text = 'Example'

    csv_fin_schema = [
        ('companyname', 'Company name', 'NVIDIA CORP'),
        ('FinancialVariable', 'Variable identifier (e.g., revt)', 'revt'),
        ('Datadate', 'Fiscal period end date', '31/12/2022'),
        ('Unit', 'Unit of measurement', 'USD million'),
        ('Value', 'Numeric value', '26914.0'),
    ]

    for i, (col, desc, example) in enumerate(csv_fin_schema, 1):
        row = table_csv_fin.rows[i].cells
        row[0].text = col
        row[1].text = desc
        row[2].text = example

    doc.add_paragraph()
    doc.add_paragraph('Table 5.8: WRDS Compustat financial data CSV schema.')

    doc.add_paragraph('Example rows:')
    doc.add_paragraph(
        'companyname,FinancialVariable,Datadate,Unit,Value\n'
        'NVIDIA CORP,revt,31/12/2022,USD million,26914.0\n'
        'NVIDIA CORP,revt,31/12/2023,USD million,60922.0',
        style='Normal'
    )

    doc.add_heading('5.6.2.4 Adding New Data Sources', level=3)

    doc.add_paragraph(
        'To integrate a new external data source, the following requirements must be satisfied:'
    )

    doc.add_paragraph('1. CSV Format: Data must be in CSV format with UTF-8 encoding.', style='List Number')
    doc.add_paragraph('2. Company Identifier: A column must uniquely identify companies.', style='List Number')
    doc.add_paragraph('3. Temporal Key: A date/year column must indicate the reporting period.', style='List Number')
    doc.add_paragraph('4. Metric Identifier: A column must contain variable names mapping to ESGMKG DatasetVariables.', style='List Number')
    doc.add_paragraph('5. Value Column: A numeric column must contain metric values.', style='List Number')

    doc.add_paragraph(
        'Once the CSV is prepared, add a new DataSource entity to the knowledge graph with SourcesFrom '
        'relationships linking relevant DatasetVariable entities. Update the DataRetrievalService to '
        'recognise and load the new data source.'
    )

    doc.add_heading('5.6.2.5 Dataset Variables and Domain Classification', level=3)

    doc.add_paragraph(
        'Dataset variables are classified by their data domain (ESG or Financial) through the hasDataDomain '
        'property. This enables correct routing of data requests and proper combination of data from '
        'heterogeneous sources for intensity calculations.'
    )

    doc.add_paragraph(
        'ESG Domain variables include: CO2DIRECTSCOPE1, CO2INDIRECTSCOPE2, SOXEMISSIONS, NOXEMISSIONS, '
        'VOCEMISSIONS, WATERCONSUMPTIONTOTAL, ENERGYUSETOTAL, HAZARDOUSWASTE (sourced from Eurofidai).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Financial Domain variables include: revt (revenue), at (total assets) (sourced from WRDS).',
        style='List Bullet'
    )

    # 5.6.3 Model Repository
    doc.add_heading('5.6.3 Model Repository', level=2)

    doc.add_paragraph(
        'The Model Repository stores computational model definitions that specify how calculated metrics are '
        'derived from their inputs. Each model is represented as an ESGMKG entity with explicit relationships '
        'to its inputs and implementation.'
    )

    doc.add_heading('5.6.3.1 Model Structure', level=3)

    doc.add_paragraph('Each model in the repository includes:')

    doc.add_paragraph('Mathematical Formula: The computation logic (e.g., "(scope1 + scope2) / revenue")', style='List Bullet')
    doc.add_paragraph('Input Requirements: References to required metrics or dataset variables via RequiresInputFrom', style='List Bullet')
    doc.add_paragraph('Implementation Reference: Link to executable code via ExecutesWith', style='List Bullet')
    doc.add_paragraph('Calculation Type: Classification of the computation (e.g., "intensity_ratio", "percentage_ratio")', style='List Bullet')

    doc.add_heading('5.6.3.2 Implemented Models', level=3)

    doc.add_paragraph('The current system includes the following calculation models:')

    # Models Table
    table_models = doc.add_table(rows=7, cols=4)
    table_models.style = 'Table Grid'

    hdr = table_models.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'Formula'
    hdr[2].text = 'Inputs'
    hdr[3].text = 'Output Unit'

    models_data = [
        ('GHGEmissionIntensityModel', '(Scope1 + Scope2) / Revenue', 'Scope1, Scope2, Revenue', 'tonnes CO2-e per M USD'),
        ('GridElectricityRateModel', '(Grid / Total) x 100', 'GridElectricity, TotalEnergy', 'Percentage (%)'),
        ('RenewableEnergyRateModel', '(Renewable / Total) x 100', 'RenewableEnergy, TotalEnergy', 'Percentage (%)'),
        ('HazardousWasteRecyclingRateModel', '(Recycled / Total) x 100', 'RecycledWaste, TotalWaste', 'Percentage (%)'),
        ('HighStressWaterConsumptionModel', '(HighStress / Total) x 100', 'HighStressWater, TotalWater', 'Percentage (%)'),
        ('HighStressWaterWithdrawalModel', '(HighStress / Total) x 100', 'HighStressWater, TotalWater', 'Percentage (%)'),
    ]

    for i, (model, formula, inputs, unit) in enumerate(models_data, 1):
        row = table_models.rows[i].cells
        row[0].text = model
        row[1].text = formula
        row[2].text = inputs
        row[3].text = unit

    doc.add_paragraph()
    doc.add_paragraph('Table 5.9: Implemented calculation models.')

    doc.add_heading('5.6.3.3 Implementation Bindings', level=3)

    doc.add_paragraph(
        'Each model is bound to a Python implementation through the ExecutesWith relationship. Implementation '
        'metadata includes:'
    )

    doc.add_paragraph('File Path: Location of the implementation script', style='List Bullet')
    doc.add_paragraph('Function Name: The callable function that executes the model', style='List Bullet')
    doc.add_paragraph('Input Parameters: Expected parameter names and types', style='List Bullet')
    doc.add_paragraph('Validation Rules: Constraints to prevent errors (e.g., division by zero checks)', style='List Bullet')

    doc.add_paragraph(
        'This separation of model definitions from implementations follows the principle of separating concerns, '
        'allowing the same computational logic to be realised through different technical means while maintaining '
        'consistent semantics.'
    )

    # 5.6.4 Internal Data Formats
    doc.add_heading('5.6.4 Internal Data Formats', level=2)

    doc.add_paragraph(
        'This section documents the internal data representations used within the system for processing, '
        'storage, and inter-service communication.'
    )

    # 5.6.4.1 Knowledge Graph Format (RDF/Turtle)
    doc.add_heading('5.6.4.1 Knowledge Graph Format (RDF/Turtle)', level=3)

    doc.add_paragraph(
        'The knowledge graph is stored in RDF Turtle (.ttl) format. Table 5.10 presents the RDF schema '
        'structure for each ESGMKG entity type.'
    )

    # RDF Schema Table
    table_rdf = doc.add_table(rows=9, cols=3)
    table_rdf.style = 'Table Grid'

    hdr = table_rdf.rows[0].cells
    hdr[0].text = 'Entity Class'
    hdr[1].text = 'Key Properties'
    hdr[2].text = 'Property Datatypes'

    rdf_schema = [
        ('esg:Industry', 'rdfs:label, esg:industryCode', 'xsd:string'),
        ('esg:ReportingFramework', 'rdfs:label, esg:frameworkVersion', 'xsd:string'),
        ('esg:Category', 'rdfs:label, esg:categoryCode', 'xsd:string'),
        ('esg:Metric', 'rdfs:label, esg:hasCalculationMethod, esg:hasDataDomain', 'xsd:string'),
        ('esg:Model', 'rdfs:label, esg:formula, esg:calculationType', 'xsd:string'),
        ('esg:Implementation', 'esg:filePath, esg:functionName', 'xsd:string'),
        ('esg:DatasetVariable', 'esg:variableName, esg:alignmentReason, esg:confidenceScore', 'xsd:string, xsd:decimal'),
        ('esg:DataSource', 'rdfs:label, esg:provider, esg:dataFormat', 'xsd:string'),
    ]

    for i, (entity, props, types) in enumerate(rdf_schema, 1):
        row = table_rdf.rows[i].cells
        row[0].text = entity
        row[1].text = props
        row[2].text = types

    doc.add_paragraph()
    doc.add_paragraph('Table 5.10: RDF schema structure for ESGMKG entities.')

    # 5.6.4.2 Service Communication Format (JSON)
    doc.add_heading('5.6.4.2 Service Communication Format (JSON)', level=3)

    doc.add_paragraph(
        'Inter-service communication and API responses use JSON format. Table 5.11 documents the key '
        'response schemas for external API consumers.'
    )

    # API Response Schema Table
    table_api_schema = doc.add_table(rows=5, cols=3)
    table_api_schema.style = 'Table Grid'

    hdr = table_api_schema.rows[0].cells
    hdr[0].text = 'Endpoint'
    hdr[1].text = 'Response Structure'
    hdr[2].text = 'Key Fields'

    api_schemas = [
        ('/api/metrics', '{"metrics": [...]}', 'id, label, unit, hasCalculationMethod, hasDataDomain'),
        ('/api/calculate', '{"results": [...]}', 'metric_id, value, unit, status, derivation_path'),
        ('/api/reports/generate', '{"report": {...}}', 'summary, metrics, audit_trail, data_sources'),
        ('/api/companies/{id}/data', '{"data": {...}}', 'company, year, variables, provenance'),
    ]

    for i, (endpoint, structure, fields) in enumerate(api_schemas, 1):
        row = table_api_schema.rows[i].cells
        row[0].text = endpoint
        row[1].text = structure
        row[2].text = fields

    doc.add_paragraph()
    doc.add_paragraph('Table 5.11: API response JSON schemas.')

    doc.add_paragraph('Example calculation response with full derivation provenance:')

    doc.add_paragraph(
        '{\n'
        '  "metric_id": "GHGEmissionIntensity",\n'
        '  "value": 42.35,\n'
        '  "unit": "tonnes CO2-e per M USD",\n'
        '  "status": "success",\n'
        '  "derivation_path": {\n'
        '    "method": "calculation_model",\n'
        '    "model": "GHGEmissionIntensityModel",\n'
        '    "inputs": [\n'
        '      {"name": "Scope1", "value": 125000, "source": "Eurofidai"},\n'
        '      {"name": "Scope2", "value": 85000, "source": "Eurofidai"},\n'
        '      {"name": "Revenue", "value": 4958, "source": "WRDS"}\n'
        '    ]\n'
        '  }\n'
        '}',
        style='Normal'
    )

    # 5.6.4.3 Runtime Data Structures
    doc.add_heading('5.6.4.3 Runtime Data Structures', level=3)

    doc.add_paragraph(
        'Table 5.12 summarises the Python data structures used internally for runtime processing.'
    )

    # Internal Data Table
    table_internal = doc.add_table(rows=5, cols=3)
    table_internal.style = 'Table Grid'

    hdr = table_internal.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Data Structure'
    hdr[2].text = 'Purpose'

    internal_data = [
        ('Knowledge Graph', 'rdflib.Graph', 'In-memory RDF store with SPARQL query interface'),
        ('Data Cache', 'pandas.DataFrame', 'Tabular ESG/Financial data indexed by (company, year)'),
        ('Calculation Results', 'Dict[str, Any]', 'Metric values with computation provenance'),
        ('Audit Trail', 'List[Dict]', 'Ordered derivation steps for traceability'),
    ]

    for i, (comp, struct, desc) in enumerate(internal_data, 1):
        row = table_internal.rows[i].cells
        row[0].text = comp
        row[1].text = struct
        row[2].text = desc

    doc.add_paragraph()
    doc.add_paragraph('Table 5.12: Runtime data structures.')

    # 5.7 User Interface Design
    doc.add_heading('5.7 User Interface Design', level=1)

    doc.add_paragraph(
        'The ESG reporting framework includes a web-based user interface that guides users through the '
        'complete ESG reporting workflow. This section presents the interface design and key screens.'
    )

    # 5.7.1 Interface Overview
    doc.add_heading('5.7.1 Interface Overview', level=2)

    doc.add_paragraph(
        'The user interface implements a step-by-step wizard pattern that guides users through the four '
        'primary use cases: framework discovery, metric exploration, metric calculation, and report generation. '
        'The interface is organised into panels that correspond to the logical workflow stages.'
    )

    doc.add_paragraph('The interface comprises the following main components:')

    doc.add_paragraph(
        'Industry Selection Panel: Allows users to select their industry sector, triggering framework discovery (UC1).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Framework Display Panel: Shows applicable reporting frameworks with descriptions and category counts.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Category and Metric Explorer: Displays categories and metrics with their properties, including '
        'calculation method indicators (UC2).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Company Selection Panel: Lists available companies for the selected industry with data availability status.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Calculation Results Panel: Displays computed metric values with derivation details and provenance (UC3).',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Report Generation Panel: Provides options for generating comprehensive reports with audit trails (UC4).',
        style='List Bullet'
    )

    # 5.7.2 Screen Layouts
    doc.add_heading('5.7.2 Screen Layouts', level=2)

    doc.add_paragraph(
        'The following figures illustrate the key screens of the user interface.'
    )

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.5: Main dashboard showing industry selection and framework discovery interface.]')

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.6: Category and metric explorer with calculation method indicators.]')

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.7: Calculation results display showing computed values with provenance details.]')

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.8: Report generation interface with audit trail preview.]')

    # 5.7.3 Visual Design Principles
    doc.add_heading('5.7.3 Visual Design Principles', level=2)

    doc.add_paragraph('The interface design follows these principles:')

    doc.add_paragraph(
        'Progressive Disclosure: Information is revealed progressively as users advance through the workflow, '
        'reducing cognitive load.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Visual Feedback: Calculation methods (direct vs. calculated) are distinguished using colour coding '
        'and icons for quick identification.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Traceability Visualisation: Derivation paths are displayed as expandable trees, allowing users to '
        'trace metrics to their sources.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Responsive Layout: The interface adapts to different screen sizes while maintaining usability.',
        style='List Bullet'
    )

    # 5.8 Sequence Diagrams
    doc.add_heading('5.8 Sequence Diagrams', level=1)

    doc.add_paragraph(
        'This section presents sequence diagrams illustrating the interactions between system components '
        'for each primary use case. These diagrams show the message flow from user action through the '
        'architectural layers to data retrieval and response.'
    )

    # 5.8.1 UC1: Report Generation Sequence
    doc.add_heading('5.8.1 UC1: ESG Report Generation', level=2)

    doc.add_paragraph(
        'Figure 5.9 illustrates the sequence of interactions for report generation. The workflow begins '
        'with framework identification and proceeds through category enumeration, metric identification, '
        'and report assembly. This use case orchestrates UC2 and UC3 for metric values.'
    )

    doc.add_paragraph('The sequence involves the following steps:')

    doc.add_paragraph('1. User selects industry in the Web Interface.', style='List Number')
    doc.add_paragraph('2. API Gateway routes request to Report Service.', style='List Number')
    doc.add_paragraph('3. Report Service invokes KG Service for framework identification (CQ1).', style='List Number')
    doc.add_paragraph('4. User selects framework; Report Service retrieves categories (CQ2).', style='List Number')
    doc.add_paragraph('5. Report Service retrieves metrics for each category (CQ3).', style='List Number')
    doc.add_paragraph('6. For each metric, Report Service invokes UC2 (Metric Calculation).', style='List Number')
    doc.add_paragraph('7. Report Service assembles report with values and provenance.', style='List Number')
    doc.add_paragraph('8. Report returned to user with full audit trail.', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.9: Sequence diagram for UC1 - ESG Report Generation showing interaction flow from User through Web Interface, API Gateway, Report Service, to KG Service, demonstrating CQ1, CQ2, CQ3.]')

    # 5.8.2 UC2: Metric Calculation Sequence
    doc.add_heading('5.8.2 UC2: Metric Calculation', level=2)

    doc.add_paragraph(
        'Figure 5.10 illustrates the metric calculation sequence. This is the most complex sequence in the '
        'system, demonstrating the distinction between direct and calculated metrics, and showing how the '
        'Calculation Service orchestrates data retrieval and model execution.'
    )

    doc.add_paragraph('Key aspects of this sequence include:')

    doc.add_paragraph(
        'Derivation Method Check: The Calculation Service first queries CQ4 to determine if the metric '
        'uses direct_measurement or calculation_model.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Direct Path: For direct metrics (hasCalculationMethod: direct_measurement), the service identifies '
        'the dataset variable via ObtainedFrom and invokes UC3 for data retrieval.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Calculated Path: For calculated metrics (hasCalculationMethod: calculation_model), the service '
        'queries CQ5 for required inputs, recursively resolves dependent metrics, identifies the model '
        'implementation via CQ6, and executes the calculation.',
        style='List Bullet'
    )
    doc.add_paragraph(
        'Cross-Domain Data: When inputs span ESG and Financial domains (hasDataDomain), the Data Retrieval '
        'Service routes requests to appropriate repositories (Eurofidai for ESG, WRDS for Financial).',
        style='List Bullet'
    )

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.10: Sequence diagram for UC2 - Metric Calculation showing interaction between Report Service, Calculation Service, KG Service, and Model Repository, demonstrating CQ4, CQ5, CQ6 and dual derivation paths.]')

    # 5.8.3 UC3: Data Access Sequence
    doc.add_heading('5.8.3 UC3: Data Access', level=2)

    doc.add_paragraph(
        'Figure 5.11 shows the data access sequence. This sequence demonstrates how the system traces '
        'metric inputs to their original data sources and retrieves values with full provenance.'
    )

    doc.add_paragraph('The sequence involves the following steps:')

    doc.add_paragraph('1. Calculation Service identifies dataset variable needed for metric.', style='List Number')
    doc.add_paragraph('2. Calculation Service invokes Data Retrieval Service.', style='List Number')
    doc.add_paragraph('3. Data Retrieval Service queries KG Service for data source (CQ7).', style='List Number')
    doc.add_paragraph('4. KG Service returns source identifier (e.g., Eurofidai, WRDS).', style='List Number')
    doc.add_paragraph('5. Data Retrieval Service routes request based on hasDataDomain.', style='List Number')
    doc.add_paragraph('6. ESG Data Repository or Financial Data Repository returns value.', style='List Number')
    doc.add_paragraph('7. Data Retrieval Service records provenance metadata.', style='List Number')
    doc.add_paragraph('8. Value with source attribution returned to Calculation Service.', style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('[FIGURE 5.11: Sequence diagram for UC3 - Data Access showing interaction between Calculation Service, Data Retrieval Service, KG Service, and ESG Data Repository, demonstrating CQ7.]')

    # 5.9 Scalability and Extensibility
    doc.add_heading('5.9 Scalability and Extensibility', level=1)

    doc.add_paragraph(
        'The framework is designed with scalability and extensibility as core architectural principles, enabling '
        'adaptation to evolving ESG reporting requirements.'
    )

    doc.add_heading('5.9.1 Framework Extensibility', level=2)

    doc.add_paragraph('New reporting frameworks (e.g., GRI, CDP, TNFD) can be added by:')

    doc.add_paragraph('Creating new ReportingFramework entities in the knowledge graph', style='List Number')
    doc.add_paragraph('Defining Category entities for each framework section', style='List Number')
    doc.add_paragraph('Mapping existing or new Metric entities to categories via ConsistsOf relationships', style='List Number')

    doc.add_paragraph(
        'No code changes are required; the system automatically discovers and exposes new frameworks through the API.'
    )

    doc.add_heading('5.9.2 Metric Extensibility', level=2)

    doc.add_paragraph('New metrics are added by:')

    doc.add_paragraph(
        'For direct metrics: Creating a Metric entity with hasCalculationMethod: "direct_measurement" and '
        'linking to a DatasetVariable via ObtainedFrom', style='List Number'
    )
    doc.add_paragraph(
        'For calculated metrics: Creating a Metric entity with hasCalculationMethod: "calculation_model", '
        'defining a Model with RequiresInputFrom relationships, and binding an Implementation', style='List Number'
    )

    doc.add_heading('5.9.3 Data Source Extensibility', level=2)

    doc.add_paragraph('New data sources are integrated by:')

    doc.add_paragraph('Creating DataSource entities with provider metadata', style='List Number')
    doc.add_paragraph('Mapping DatasetVariable entities via SourcesFrom relationships', style='List Number')
    doc.add_paragraph('Implementing data retrieval adapters in the Data Retrieval Service', style='List Number')

    doc.add_heading('5.9.4 Industry Extensibility', level=2)

    doc.add_paragraph(
        'The SASB framework covers 77 industries across 11 sectors. Adding a new industry requires:'
    )

    doc.add_paragraph('Creating an Industry entity', style='List Number')
    doc.add_paragraph('Linking to appropriate ReportingFramework entities via ReportsUsing', style='List Number')
    doc.add_paragraph('Populating industry-specific metrics and data mappings', style='List Number')

    doc.add_paragraph(
        'This modular design ensures that the system can scale to accommodate the full breadth of ESG reporting '
        'requirements across industries, frameworks, and data providers without architectural changes.'
    )

    # 5.10 Summary
    doc.add_heading('5.10 Summary', level=1)

    doc.add_paragraph(
        'This chapter presented the implementation of an ESG reporting framework that operationalises the '
        'ESGMKG ontology. The four-layer architecture separates presentation, API, service, and data concerns, '
        'enabling modularity and maintainability. Four microservices implement the core functionality: the '
        'Knowledge Graph Service for semantic navigation, the Data Retrieval Service for data access across '
        'multiple sources, the Calculation Service for metric derivation using both direct and calculated paths, '
        'and the Report Service for comprehensive report generation.'
    )

    doc.add_paragraph(
        'Key implementation features include:'
    )

    doc.add_paragraph(
        'Full implementation of all seven competency questions through SPARQL queries', style='List Bullet'
    )
    doc.add_paragraph(
        'Support for both direct measurement retrieval and computational model execution', style='List Bullet'
    )
    doc.add_paragraph(
        'Cross-domain data integration using hasDataDomain to combine ESG and Financial data', style='List Bullet'
    )
    doc.add_paragraph(
        'Complete audit trail from reported values to underlying data sources', style='List Bullet'
    )
    doc.add_paragraph(
        'Extensible architecture supporting new frameworks, metrics, and data sources', style='List Bullet'
    )

    doc.add_paragraph(
        'The implementation validates the ESGMKG ontology design by demonstrating that the conceptual model '
        'can be effectively operationalised to support practical ESG reporting workflows. The separation of '
        'ontological knowledge from procedural logic enables flexibility and adaptability as ESG reporting '
        'requirements evolve.'
    )

    # Save document
    doc.save('/Users/mingqin/Downloads/esg-knowledge-graph-demo/Demo corresponding acadedmic paper/Chapter5_ESG_Reporting_Framework_Implementation.docx')
    print("Chapter 5 saved successfully!")


if __name__ == '__main__':
    create_chapter4()
    create_chapter5()
    print("\nBoth chapters generated successfully!")
    print("Files saved to:")
    print("  - Chapter4_ESG_Metric_Knowledge_Graph.docx")
    print("  - Chapter5_ESG_Reporting_Framework_Implementation.docx")
